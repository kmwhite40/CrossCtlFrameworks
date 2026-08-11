"""Acceptance -- the human gate that projects a proposal into the existing
``AssessmentControlResult`` shape the SAR generator and POA&M path already read.
"""

from __future__ import annotations

import io
from typing import Any

import pytest
from docx import Document
from sqlalchemy import delete, select

from ccf.ai_actions.provenance import record_ai_run
from ccf.assessment.engine import service
from ccf.assessment.engine.evaluate import ObjectiveEvaluation
from ccf.assessment.engine.service import (
    AcceptanceRefused,
    accept_control_proposal,
    check_staleness,
    evaluate_control_proposal,
    open_control_proposal,
)
from ccf.assessment.sar import generate_sar_docx
from ccf.assessment.seed import result_to_dict, summarize_results
from ccf.db import session_scope
from ccf.models import Assessment, AssessmentControlResult, Control, Organization, System
from ccf.models_ai_actions import AiActionRun
from ccf.models_assessment_engine import AssessmentControlProposal, AssessmentObjectiveProposal

pytestmark = pytest.mark.usefixtures("fresh_engine")

_SEQ = "ZQ-90"


@pytest.fixture(autouse=True)
async def _catalog_rows():
    """Seed one addressable control plus three sub-clause objective rows.

    Mirrors tests/test_assessment_service.py's fixture shape: a parent row with
    a bare "Determine if:" header, and three sub-clause rows carrying the actual
    objective text with control_name NULL.
    """
    async with session_scope() as s:
        await s.execute(delete(Control).where(Control.sequence_control == _SEQ))
        s.add(
            Control(
                identifier=_SEQ,
                sequence_control=_SEQ,
                control_name="Test Policy And Procedures",
                assessment_objective="Determine if:",
                source_row=1,
            )
        )
        s.add(
            Control(
                identifier=f"{_SEQ}-ao1",
                sequence_control=_SEQ,
                ap_acronym="ZQ-90a",
                assessment_objective="personnel to whom the policy is disseminated are defined;",
                source_row=2,
            )
        )
        s.add(
            Control(
                identifier=f"{_SEQ}-ao2",
                sequence_control=_SEQ,
                ap_acronym="ZQ-90b",
                assessment_objective="an official to manage the policy is defined;",
                source_row=3,
            )
        )
        s.add(
            Control(
                identifier=f"{_SEQ}-ao3",
                sequence_control=_SEQ,
                ap_acronym="ZQ-90c",
                assessment_objective="the review frequency is defined;",
                source_row=4,
            )
        )
    yield
    async with session_scope() as s:
        await s.execute(delete(Control).where(Control.sequence_control == _SEQ))


async def _assessment(name: str) -> tuple[int, int]:
    """Create an org + system + assessment. Returns (org_id, assessment_id)."""
    async with session_scope() as s:
        org = Organization(name=name)
        s.add(org)
        await s.flush()
        system = System(organization_id=org.id, name=f"{name}-system")
        s.add(system)
        await s.flush()
        assessment = Assessment(system_id=system.id, name=f"{name}-assessment", kind="self")
        s.add(assessment)
        await s.flush()
        return int(org.id), int(assessment.id)


def _fake_evaluate_always(verdict: str = "satisfied"):
    async def _fake(session: Any, **kwargs: Any) -> ObjectiveEvaluation:
        return ObjectiveEvaluation(verdict=verdict, rationale="ok", confidence=0.5)

    return _fake


async def _evaluated_proposal(
    name: str, monkeypatch: pytest.MonkeyPatch, verdict: str = "satisfied"
) -> int:
    """Open + evaluate a proposal against the ZQ-90 fixture. Returns its id."""
    monkeypatch.setattr(service, "evaluate_objective", _fake_evaluate_always(verdict))
    _, assessment_id = await _assessment(name)
    async with session_scope() as s:
        proposal = await open_control_proposal(
            s, assessment_id=assessment_id, control_identifier="ZQ-90"
        )
        proposal = await evaluate_control_proposal(s, proposal)
        return int(proposal.id)


_SEQ2 = "ZQ-91"


@pytest.fixture
async def _second_catalog_rows():
    """A second addressable control, for tests that need two proposals in the
    same assessment. Separate from ``_catalog_rows`` (autouse, ZQ-90 only) so
    every other test in this module keeps its single-control shape.
    """
    async with session_scope() as s:
        await s.execute(delete(Control).where(Control.sequence_control == _SEQ2))
        s.add(
            Control(
                identifier=_SEQ2,
                sequence_control=_SEQ2,
                control_name="Test Policy And Procedures Two",
                assessment_objective="Determine if:",
                source_row=1,
            )
        )
        s.add(
            Control(
                identifier=f"{_SEQ2}-ao1",
                sequence_control=_SEQ2,
                ap_acronym="ZQ-91a",
                assessment_objective="a second control's objective is defined;",
                source_row=2,
            )
        )
    yield
    async with session_scope() as s:
        await s.execute(delete(Control).where(Control.sequence_control == _SEQ2))


def _fake_evaluate_with_provenance(
    verdict: str = "satisfied", *, null_run_label: str | None = None
):
    """Like ``_fake_evaluate_always``, but calls ``record_ai_run`` for real --
    the same way ``evaluate_objective`` itself does -- so acceptance-stamping
    tests exercise the actual ``AiActionRun`` FK rather than a fabricated
    integer. ``null_run_label`` simulates a provenance-recording failure on
    one objective (its ``ai_action_run_id`` stays NULL, exactly what
    ``evaluate_objective`` returns when ``record_ai_run`` itself fails),
    leaving every other objective with a real linked run.
    """

    async def _fake(
        session: Any, *, objective: Any, org_id: int, **kwargs: Any
    ) -> ObjectiveEvaluation:
        if objective.label == null_run_label:
            return ObjectiveEvaluation(verdict=verdict, rationale="ok", confidence=0.5)
        ai_run = await record_ai_run(
            session,
            action_key="evaluate_assessment_objective",
            entity_type="assessment_objective",
            entity_id=objective.label,
            organization_id=org_id,
            provider="anthropic",
            model="fake-model",
            prompt=f"evaluate {objective.label}",
            output={"verdict": verdict},
            citations=[],
        )
        assert ai_run is not None
        return ObjectiveEvaluation(
            verdict=verdict, rationale="ok", confidence=0.5, ai_action_run_id=ai_run.id
        )

    return _fake


async def _evaluated_proposal_with_runs(
    name: str,
    monkeypatch: pytest.MonkeyPatch,
    verdict: str = "satisfied",
    *,
    control_identifier: str = "ZQ-90",
    null_run_label: str | None = None,
) -> int:
    """Like ``_evaluated_proposal``, but every objective's evaluation records a
    real ``AiActionRun`` (unless ``null_run_label`` names one to skip).
    """
    fake = _fake_evaluate_with_provenance(verdict, null_run_label=null_run_label)
    monkeypatch.setattr(service, "evaluate_objective", fake)
    _, assessment_id = await _assessment(name)
    async with session_scope() as s:
        proposal = await open_control_proposal(
            s, assessment_id=assessment_id, control_identifier=control_identifier
        )
        proposal = await evaluate_control_proposal(s, proposal)
        return int(proposal.id)


async def _run_ids_for_proposal(proposal_id: int) -> list[int]:
    """The non-NULL ``ai_action_run_id`` values linked to one proposal's objectives."""
    async with session_scope() as s:
        ids = (
            await s.execute(
                select(AssessmentObjectiveProposal.ai_action_run_id).where(
                    AssessmentObjectiveProposal.control_proposal_id == proposal_id
                )
            )
        ).scalars().all()
    return [i for i in ids if i is not None]


async def _runs(run_ids: list[int]) -> list[AiActionRun]:
    async with session_scope() as s:
        result = await s.execute(select(AiActionRun).where(AiActionRun.id.in_(run_ids)))
        return list(result.scalars())


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


async def test_acceptance_projects_objectives_into_the_existing_shape(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """objective_findings must match what ccf.assessment.sar already renders."""
    proposal_id = await _evaluated_proposal("acc-shape", monkeypatch)
    async with session_scope() as s:
        result = await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")
        objective_findings = list(result.objective_findings)

    assert objective_findings == [
        {
            "label": "ZQ-90a",
            "text": "personnel to whom the policy is disseminated are defined;",
            "finding": "satisfied",
        },
        {
            "label": "ZQ-90b",
            "text": "an official to manage the policy is defined;",
            "finding": "satisfied",
        },
        {
            "label": "ZQ-90c",
            "text": "the review frequency is defined;",
            "finding": "satisfied",
        },
    ]


async def test_acceptance_sets_the_control_finding(monkeypatch: pytest.MonkeyPatch) -> None:
    proposal_id = await _evaluated_proposal("acc-finding", monkeypatch, verdict="satisfied")
    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        result = await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")
        assert result.finding == proposal.proposed_finding == "satisfied"


async def test_acceptance_records_actor_and_timestamp(monkeypatch: pytest.MonkeyPatch) -> None:
    proposal_id = await _evaluated_proposal("acc-actor", monkeypatch)
    async with session_scope() as s:
        await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")

    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        assert proposal.state == "accepted"
        assert proposal.accepted_by == "assessor@example.com"
        assert proposal.accepted_at is not None


async def test_an_insufficient_evidence_proposal_cannot_be_accepted(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """'The engine could not tell' must never become a finding."""
    proposal_id = await _evaluated_proposal(
        "acc-insufficient", monkeypatch, verdict="insufficient_evidence"
    )
    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        assert proposal.proposed_finding == "insufficient_evidence"
        assessment_id = proposal.assessment_id
        with pytest.raises(AcceptanceRefused):
            await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")

    async with session_scope() as s:
        rows = (
            await s.execute(
                select(AssessmentControlResult).where(
                    AssessmentControlResult.assessment_id == assessment_id
                )
            )
        ).scalars().all()
    assert rows == []


async def test_a_stale_proposal_cannot_be_accepted(monkeypatch: pytest.MonkeyPatch) -> None:
    proposal_id = await _evaluated_proposal("acc-stale", monkeypatch)

    async with session_scope() as s:
        control = (
            await s.execute(select(Control).where(Control.identifier == f"{_SEQ}-ao1"))
        ).scalar_one()
        control.assessment_objective = "a completely reworded objective statement;"

    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        is_stale = await check_staleness(s, proposal)
        assert is_stale is True

        # Match on wording unique to the stale-specific branch, not merely the
        # "not complete" superset -- "stale" != "complete" would also refuse,
        # so an exact-exception-type check alone cannot tell them apart.
        with pytest.raises(AcceptanceRefused, match="catalog objective text changed"):
            await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")


async def test_accept_refuses_a_reworded_objective_without_a_manual_staleness_check(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """I1: check_staleness was previously only ever called by tests, never by
    accept_control_proposal itself -- production code had no path that ever
    marked a proposal stale, so the staleness guard right below in this same
    function was unreachable in practice, and objective_text_sha256 bought
    nothing. Unlike test_a_stale_proposal_cannot_be_accepted above, this test
    deliberately does NOT call check_staleness before accepting -- proving
    accept_control_proposal now detects the reword on its own.
    """
    proposal_id = await _evaluated_proposal("acc-auto-stale", monkeypatch)

    async with session_scope() as s:
        control = (
            await s.execute(select(Control).where(Control.identifier == f"{_SEQ}-ao1"))
        ).scalar_one()
        control.assessment_objective = "a completely reworded objective statement;"

    async with session_scope() as s:
        with pytest.raises(AcceptanceRefused, match="catalog objective text changed"):
            await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")

    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        assert proposal.state == "stale"

        results = (
            await s.execute(
                select(AssessmentControlResult).where(
                    AssessmentControlResult.assessment_id == proposal.assessment_id
                )
            )
        ).scalars().all()
        assert results == [], "a stale proposal must not project into AssessmentControlResult"


async def test_accepting_an_already_accepted_proposal_is_refused(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A second immediate accept, with no re-evaluation in between, must still
    be refused. The proposal's ``proposed_finding`` is still set from the
    first evaluation and is not ``insufficient_evidence``, and its state is
    ``"stale"``'s sibling ``"accepted"`` rather than ``"complete"`` -- so only
    the ``state != "complete"`` branch catches this; every other guard passes
    it through. Task 11 wires this path to a POST endpoint a client can call
    twice, so a repeated request must not silently re-accept.
    """
    proposal_id = await _evaluated_proposal("acc-double-immediate", monkeypatch)
    async with session_scope() as s:
        await accept_control_proposal(s, proposal_id, accepted_by="first@example.com")

    async with session_scope() as s:
        with pytest.raises(AcceptanceRefused, match="not complete"):
            await accept_control_proposal(s, proposal_id, accepted_by="second@example.com")


async def test_accepting_twice_does_not_duplicate_the_result_row(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """uq_assess_ctrl on (assessment_id, control_id) -- second call updates."""
    monkeypatch.setattr(service, "evaluate_objective", _fake_evaluate_always("satisfied"))
    _, assessment_id = await _assessment("acc-twice")
    async with session_scope() as s:
        proposal = await open_control_proposal(
            s, assessment_id=assessment_id, control_identifier="ZQ-90"
        )
        proposal = await evaluate_control_proposal(s, proposal)
        proposal_id = int(proposal.id)
        result = await accept_control_proposal(s, proposal_id, accepted_by="first@example.com")
        assert result.finding == "satisfied"

    # Re-evaluate under a different verdict and accept again -- must update the
    # one existing row, not violate uq_assess_ctrl with a second insert.
    monkeypatch.setattr(service, "evaluate_objective", _fake_evaluate_always("not_satisfied"))
    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        proposal = await evaluate_control_proposal(s, proposal)
        result = await accept_control_proposal(s, proposal_id, accepted_by="second@example.com")
        assert result.finding == "other_than_satisfied"

    async with session_scope() as s:
        rows = (
            await s.execute(
                select(AssessmentControlResult).where(
                    AssessmentControlResult.assessment_id == assessment_id
                )
            )
        ).scalars().all()
    assert len(rows) == 1
    assert rows[0].finding == "other_than_satisfied"
    assert rows[0].reviewer == "second@example.com"


async def test_the_generated_sar_renders_the_projected_objectives(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The whole point of reusing AssessmentControlResult."""
    proposal_id = await _evaluated_proposal("acc-sar", monkeypatch, verdict="satisfied")
    async with session_scope() as s:
        result = await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")
        result_dict = result_to_dict(result)
        summary = summarize_results([result])

    meta = {
        "customer_name": "Acme Corp",
        "system_name": "Acme System",
        "assessment_name": "acc-sar-assessment",
        "kind": "self",
        "assessor": "assessor@example.com",
        "period": "2026-01-01 to ongoing",
        "date": "08/11/2026",
    }
    data = generate_sar_docx(meta, summary, [result_dict])
    text = _docx_text(data)

    assert "ZQ-90a" in text
    assert "personnel to whom the policy is disseminated are defined;" in text


async def test_before_acceptance_no_linked_run_is_stamped(monkeypatch: pytest.MonkeyPatch) -> None:
    """reviewer, disposition, decided_at all NULL; mutation_applied False."""
    proposal_id = await _evaluated_proposal_with_runs("acc-run-unstamped-before", monkeypatch)
    run_ids = await _run_ids_for_proposal(proposal_id)
    assert len(run_ids) == 3, "every objective in the ZQ-90 fixture should have recorded a run"

    for run in await _runs(run_ids):
        assert run.reviewer is None
        assert run.disposition is None
        assert run.decided_at is None
        assert run.mutation_applied is False


async def test_acceptance_stamps_every_linked_run(monkeypatch: pytest.MonkeyPatch) -> None:
    """Each run gets reviewer == accepted_by, disposition 'accepted',
    a decided_at, and mutation_applied True.
    """
    proposal_id = await _evaluated_proposal_with_runs("acc-run-stamp-all", monkeypatch)
    run_ids = await _run_ids_for_proposal(proposal_id)
    assert len(run_ids) == 3

    async with session_scope() as s:
        await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")

    runs = await _runs(run_ids)
    assert len(runs) == 3
    for run in runs:
        assert run.reviewer == "assessor@example.com"
        assert run.disposition == "accepted"
        assert run.decided_at is not None
        assert run.mutation_applied is True


async def test_a_proposal_with_a_null_run_id_still_accepts(monkeypatch: pytest.MonkeyPatch) -> None:
    """A provenance failure must not block acceptance."""
    proposal_id = await _evaluated_proposal_with_runs(
        "acc-run-null-id", monkeypatch, null_run_label="ZQ-90b"
    )
    all_run_ids = await _run_ids_for_proposal(proposal_id)
    assert len(all_run_ids) == 2, "ZQ-90b's provenance failure leaves only two linked runs"

    async with session_scope() as s:
        result = await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")
    assert result.finding == "satisfied"

    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.id == proposal_id
                )
            )
        ).scalar_one()
        assert proposal.state == "accepted"

        by_label = {
            o.label: o
            for o in (
                await s.execute(
                    select(AssessmentObjectiveProposal).where(
                        AssessmentObjectiveProposal.control_proposal_id == proposal_id
                    )
                )
            )
            .scalars()
            .all()
        }
    assert by_label["ZQ-90b"].ai_action_run_id is None

    for run in await _runs(all_run_ids):
        assert run.reviewer == "assessor@example.com"
        assert run.disposition == "accepted"
        assert run.decided_at is not None
        assert run.mutation_applied is True


async def test_a_refused_acceptance_stamps_nothing(monkeypatch: pytest.MonkeyPatch) -> None:
    """An insufficient_evidence proposal is refused and leaves runs unstamped."""
    proposal_id = await _evaluated_proposal_with_runs(
        "acc-run-refused", monkeypatch, verdict="insufficient_evidence"
    )
    run_ids = await _run_ids_for_proposal(proposal_id)
    assert len(run_ids) == 3

    async with session_scope() as s:
        with pytest.raises(AcceptanceRefused):
            await accept_control_proposal(s, proposal_id, accepted_by="assessor@example.com")

    for run in await _runs(run_ids):
        assert run.reviewer is None
        assert run.disposition is None
        assert run.decided_at is None
        assert run.mutation_applied is False


async def test_accepting_one_proposal_does_not_stamp_a_different_proposals_runs(
    monkeypatch: pytest.MonkeyPatch, _second_catalog_rows: None
) -> None:
    """Two control proposals in the same assessment, each with a linked run --
    accepting one must not touch the other's.
    """
    monkeypatch.setattr(service, "evaluate_objective", _fake_evaluate_with_provenance())
    _, assessment_id = await _assessment("acc-run-cross-proposal")
    async with session_scope() as s:
        proposal_a = await open_control_proposal(
            s, assessment_id=assessment_id, control_identifier="ZQ-90"
        )
        proposal_a = await evaluate_control_proposal(s, proposal_a)
        proposal_a_id = int(proposal_a.id)

        proposal_b = await open_control_proposal(
            s, assessment_id=assessment_id, control_identifier="ZQ-91"
        )
        proposal_b = await evaluate_control_proposal(s, proposal_b)
        proposal_b_id = int(proposal_b.id)

    run_ids_a = await _run_ids_for_proposal(proposal_a_id)
    run_ids_b = await _run_ids_for_proposal(proposal_b_id)
    assert len(run_ids_a) == 3
    assert len(run_ids_b) == 1

    async with session_scope() as s:
        await accept_control_proposal(s, proposal_a_id, accepted_by="assessor@example.com")

    for run in await _runs(run_ids_a):
        assert run.reviewer == "assessor@example.com"
        assert run.disposition == "accepted"
        assert run.mutation_applied is True

    for run in await _runs(run_ids_b):
        assert run.reviewer is None
        assert run.disposition is None
        assert run.decided_at is None
        assert run.mutation_applied is False
