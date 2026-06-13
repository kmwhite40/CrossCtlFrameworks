"""Unit tests for the SPRS scoring engine, parser, and SSP generator (no DB)."""

from __future__ import annotations

import io

from docx import Document

from ccf.scoring.engine import (
    MAX_SPRS_SCORE,
    SPRS_FLOOR,
    deduction_for,
    score_system,
)
from ccf.scoring.parser import load_seed, split_objectives
from ccf.ssp.generator import generate_ssp_docx


def test_deduction_weights() -> None:
    # Full miss subtracts the point value.
    assert deduction_for("5", "not_implemented") == 5
    assert deduction_for("3", "not_implemented") == 3
    assert deduction_for("1", "not_implemented") == 1
    # Met-like states never deduct.
    for state in ("implemented", "inherited", "not_applicable"):
        assert deduction_for("5", state) == 0
    # Partial credit only on the 3/5 rows; others lose the full value.
    assert deduction_for("3/5", "partial") == 3
    assert deduction_for("3/5", "not_implemented") == 5
    assert deduction_for("5", "partial") == 5
    # Special (SSP) carries no numeric weight.
    assert deduction_for("Special", "not_implemented") == 0


def test_seed_has_110_controls() -> None:
    seed = load_seed()
    assert len(seed) == 110
    assert {r["point_value"] for r in seed} <= {"1", "3", "3/5", "5", "Special"}
    assert all(r["control_id"] for r in seed)


def test_split_objectives_parts() -> None:
    parts = split_objectives("[a] foo is identified;\n[b] bar is controlled.")
    assert [p["label"] for p in parts] == ["a", "b"]
    assert parts[0]["text"] == "foo is identified"
    # No markers → a single unlabeled part.
    assert split_objectives("plain text")[0]["label"] == ""


def test_score_system_extremes_and_floor() -> None:
    seed = load_seed()
    # Everything implemented → perfect 110.
    perfect = score_system(seed, {r["control_id"]: "implemented" for r in seed})
    assert perfect.score == MAX_SPRS_SCORE
    assert perfect.met_controls == 110
    # Nothing assessed → maximum deduction, clamped at the SPRS floor.
    worst = score_system(seed, {})
    assert worst.score == SPRS_FLOOR
    assert worst.deductions_total == 313
    assert worst.ssp_present is False


def test_score_recomputes_per_control() -> None:
    seed = load_seed()
    states = {r["control_id"]: "implemented" for r in seed}
    five = next(r["control_id"] for r in seed if r["point_value"] == "5")
    states[five] = "not_implemented"
    summary = score_system(seed, states)
    assert summary.score == MAX_SPRS_SCORE - 5
    assert summary.by_domain  # populated breakdown


def test_generate_ssp_docx_roundtrips() -> None:
    project = {"customer_name": "Acme", "system_name": "Enclave", "version": "0.1"}
    entries = [
        {
            "control_id": "AC.L2-3.1.1",
            "nist_id": "3.1.1",
            "domain": "AC",
            "title": "Authorized Access Control",
            "responsible_role": "Access Control Lead",
            "implementation_status": ["Implemented"],
            "control_origination": ["Shared"],
            "part_narratives": [{"label": "a", "text": "users are identified."}],
        }
    ]
    data = generate_ssp_docx(project, entries)
    doc = Document(io.BytesIO(data))
    para_text = "\n".join(p.text for p in doc.paragraphs)
    assert "System Security Plan" in para_text
    cell_text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "AC.L2-3.1.1" in cell_text
    assert "users are identified" in cell_text
