"""FR-01: the SSP generator must honestly label its true framework —
CMMC Level 2 / NIST SP 800-171 Rev.2 — and never claim FedRAMP / NIST 800-53
content that the pipeline does not actually produce.

These tests exercise the real production path: the actual ``.docx`` bytes
returned by ``generate_ssp_docx`` (parsed back with python-docx) and the
actual ``QUESTIONNAIRE`` options object served to the intake UI/API — not a
synthetic stand-in for either.
"""

from __future__ import annotations

import io

from docx import Document

from ccf.governance.automation import QUESTIONNAIRE
from ccf.ssp.generator import generate_ssp_docx

PROJECT = {
    "title": "System Security Plan (SSP)",
    "customer_name": "Acme Federal LLC",
    "system_name": "Acme Enclave",
    "platform": "Microsoft 365 GCC High",
    "version": "0.1",
    "document_date": "01/01/2026",
    "prepared_by": "Jane Doe",
}

ENTRIES = [
    {
        "control_id": "AC.L2-3.1.1",
        "nist_id": "3.1.1",
        "domain": "AC",
        "title": "Authorized Access Control",
        "responsible_role": "Access Control Lead",
        "implementation_status": ["Implemented"],
        "control_origination": ["Organization System Specific"],
        "part_narratives": [{"label": "a", "text": "Access is limited to authorized users."}],
    }
]


def _all_text(docx_bytes: bytes) -> str:
    """Flatten every paragraph and table-cell run into one searchable string."""
    doc = Document(io.BytesIO(docx_bytes))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def test_cover_names_the_true_framework_not_fedramp() -> None:
    text = _all_text(generate_ssp_docx(PROJECT, ENTRIES))

    # Must name the real basis of the document.
    assert "CMMC" in text
    assert "NIST SP 800-171 Rev. 2" in text

    # Must not claim this is a FedRAMP SSP or 800-53 content.
    assert "FedRAMP SSP" not in text
    assert "800-53" not in text


def test_framework_picker_does_not_offer_fedramp_or_800_53() -> None:
    frameworks_question = next(q for q in QUESTIONNAIRE if q["id"] == "frameworks")
    options = frameworks_question["options"]

    assert "FedRAMP" not in options
    assert "NIST_800_53" not in options
    # The framework the generator actually produces must still be selectable.
    assert "CMMC_L2" in options
    assert "NIST_800_171" in options
