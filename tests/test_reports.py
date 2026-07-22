"""Custom report builder export formats (CSV / XLSX / DOCX / JSON).

Also covers CISO-09 / CISO-10 leadership decision-support polish:

- ``ccf.analytics.posture.org_summary`` surfaces the worst/min SPRS-scored
  system alongside the average, so a failing system can't hide behind a
  healthy mean (CISO-09).
- the exported compliance report carries a POA&M/risk posture summary
  reconciled — same numbers, not a re-derivation — to ``org_summary`` for the
  same scope, plus an AI/last-editor provenance flag on rows whose
  implementation status is backed by a still-AI-drafted SSP narrative,
  reusing the CISO-02 ``DRAFT_PREFIX`` signal (CISO-10).
"""

from __future__ import annotations

import io
from datetime import date

import openpyxl
import pytest
from alembic import command
from alembic.config import Config
from docx import Document
from httpx import ASGITransport, AsyncClient
from sqlalchemy import select

from ccf.analytics import posture
from ccf.api.main import create_app
from ccf.config import get_settings
from ccf.db import session_scope
from ccf.models import (
    POAM,
    Control,
    ControlImplementation,
    Organization,
    ScoringControl,
    ScoringStatus,
    SSPControlEntry,
    SSPProject,
    System,
)
from ccf.reporting import report_to_docx, report_to_xlsx
from ccf.scoring.seed import seed_scoring_controls
from ccf.ssp.statements import DRAFT_PREFIX

pytestmark = pytest.mark.usefixtures("fresh_engine")

_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

TODAY = date(2026, 7, 21)


@pytest.fixture(scope="module", autouse=True)
def _migrate() -> None:
    cfg = Config("alembic.ini")
    cfg.set_main_option("sqlalchemy.url", str(get_settings().database_url_sync))
    command.upgrade(cfg, "head")


def _client() -> AsyncClient:
    return AsyncClient(transport=ASGITransport(app=create_app()), base_url="http://t")


def test_renderers_emit_real_office_files() -> None:
    summary = {"organization": "Acme", "baseline": "mod", "total_rows": 1}
    rows = [
        {"identifier": "AC-1", "family": "AC", "control_name": "Policy", "baseline_mod": True},
    ]
    xlsx = report_to_xlsx(summary, rows)
    docx = report_to_docx(summary, rows)
    assert xlsx[:2] == b"PK" and len(xlsx) > 0
    assert docx[:2] == b"PK" and len(docx) > 0


@pytest.mark.asyncio
async def test_build_endpoint_serves_each_format() -> None:
    async with _client() as c:
        j = await c.get("/api/reports/build", params={"baseline": "mod", "fmt": "json"})
        assert j.status_code == 200
        assert {"summary", "risk_summary", "rows"} <= j.json().keys()

        xlsx = await c.get(
            "/api/reports/build", params={"baseline": "mod", "fmt": "xlsx", "filename": "audit pkg"}
        )
        assert xlsx.status_code == 200
        assert xlsx.headers["content-type"] == _XLSX
        assert "audit_pkg.xlsx" in xlsx.headers["content-disposition"]
        assert xlsx.content[:2] == b"PK"

        docx = await c.get("/api/reports/build", params={"baseline": "mod", "fmt": "docx"})
        assert docx.status_code == 200
        assert docx.headers["content-type"] == _DOCX
        assert docx.content[:2] == b"PK"


# --- CISO-09: org_summary surfaces the worst/min SPRS-scored system --------


async def _seeded_scoring_controls() -> list[tuple[int, str]]:
    """Load the real CMMC L2 scoring matrix; return (scoring_control_id, control_id) pairs."""
    async with session_scope() as s:
        await seed_scoring_controls(s)
        rows = (await s.execute(select(ScoringControl.id, ScoringControl.control_id))).all()
        return [(sid, cid) for sid, cid in rows]


@pytest.mark.asyncio
async def test_org_summary_exposes_worst_system_sprs() -> None:
    async with session_scope() as s:
        org = Organization(name="SprsWorstOrg")
        s.add(org)
        await s.flush()
        good = System(organization_id=org.id, name="GoodSys", baseline="moderate")
        mid = System(organization_id=org.id, name="MidSys", baseline="moderate")
        bad = System(organization_id=org.id, name="BadSys", baseline="moderate")
        s.add_all([good, mid, bad])
        await s.flush()
        org_id, good_id, mid_id, bad_id = org.id, good.id, mid.id, bad.id

    scoring_controls = await _seeded_scoring_controls()
    assert scoring_controls, "the CMMC L2 seed must produce scoring controls"

    async with session_scope() as s:
        # GoodSys: fully implemented -> perfect (or near-perfect) score.
        for scoring_id, _ in scoring_controls:
            s.add(
                ScoringStatus(system_id=good_id, scoring_control_id=scoring_id, state="implemented")
            )
        # MidSys: fully implemented except one control -> a little below Good.
        for i, (scoring_id, _) in enumerate(scoring_controls):
            state = "partial" if i == 0 else "implemented"
            s.add(ScoringStatus(system_id=mid_id, scoring_control_id=scoring_id, state=state))
        # BadSys: only a handful of controls assessed at all, and all missed —
        # deeply negative, unambiguously the weakest assessed system.
        for scoring_id, _ in scoring_controls[:3]:
            s.add(
                ScoringStatus(
                    system_id=bad_id, scoring_control_id=scoring_id, state="not_implemented"
                )
            )

    async with session_scope() as s:
        cards = await posture.systems_scorecard(s, today=TODAY, org_id=org_id)
        summary = await posture.org_summary(s, today=TODAY, org_id=org_id)

    scored = {c["system_id"]: c for c in cards if c["controls_assessed"] > 0}
    assert set(scored) == {good_id, mid_id, bad_id}
    expected_worst = min(scored.values(), key=lambda c: c["sprs_score"])

    # BadSys really is the weakest — sanity-check the setup, not just the API.
    assert expected_worst["system_id"] == bad_id

    assert summary["min_sprs_score"] == expected_worst["sprs_score"]
    assert summary["worst_system"] is not None
    assert summary["worst_system"]["system_id"] == bad_id
    assert summary["worst_system"]["name"] == "BadSys"
    # The mean masks exactly what this feature is meant to unmask: the worst
    # system scores well below the average of the portfolio.
    assert summary["min_sprs_score"] < summary["avg_sprs_score"]


@pytest.mark.asyncio
async def test_org_summary_worst_system_none_when_no_system_scored() -> None:
    async with session_scope() as s:
        org = Organization(name="NoScoreOrg")
        s.add(org)
        await s.flush()
        sysm = System(organization_id=org.id, name="UnscoredSys", baseline="moderate")
        s.add(sysm)
        await s.flush()
        org_id = org.id

    async with session_scope() as s:
        summary = await posture.org_summary(s, today=TODAY, org_id=org_id)

    assert summary["systems_scored"] == 0
    assert summary["min_sprs_score"] is None
    assert summary["worst_system"] is None


# --- CISO-10: export risk/POA&M summary reconciles to org_summary ----------


@pytest.mark.asyncio
async def test_report_risk_summary_reconciles_to_org_summary() -> None:
    async with session_scope() as s:
        org = Organization(name="ReportRiskOrg")
        s.add(org)
        await s.flush()
        sysm = System(organization_id=org.id, name="ReportRiskSys", baseline="moderate")
        s.add(sysm)
        await s.flush()
        s.add_all(
            [
                POAM(
                    system_id=sysm.id,
                    title="Overdue weakness",
                    severity="high",
                    status="open",
                    identified_on=date(2020, 1, 1),
                    due_on=date(2020, 2, 1),
                ),
                POAM(
                    system_id=sysm.id,
                    title="On-track weakness",
                    severity="low",
                    status="open",
                    identified_on=date(2020, 1, 1),
                    due_on=date(2999, 1, 1),
                ),
                POAM(
                    system_id=sysm.id,
                    title="Accepted residual risk",
                    severity="high",
                    status="risk_accepted",
                    identified_on=date(2020, 1, 1),
                ),
            ]
        )
        org_id = org.id

    async with _client() as c:
        j = await c.get(
            "/api/reports/build", params={"organization_id": org_id, "fmt": "json"}
        )
        assert j.status_code == 200
        body = j.json()
        risk_summary = body["risk_summary"]

        async with session_scope() as s:
            expected = await posture.org_summary(s, today=date.today(), org_id=org_id)

        assert risk_summary["open_poams"] == expected["open_poams"] == 2
        assert risk_summary["overdue_poams"] == expected["overdue_poams"] == 1
        assert risk_summary["accepted_poams"] == expected["accepted_poams"] == 1
        assert risk_summary["systems_total"] == expected["systems_total"]
        assert risk_summary["systems_scored"] == expected["systems_scored"]
        assert risk_summary["avg_sprs_score"] == expected["avg_sprs_score"]
        assert risk_summary["min_sprs_score"] == expected["min_sprs_score"]
        assert risk_summary["worst_system"] == expected["worst_system"]
        assert risk_summary["risks_by_status"] == expected["risks_by_status"]

        # xlsx: same numbers land on the Summary sheet.
        xlsx_resp = await c.get(
            "/api/reports/build", params={"organization_id": org_id, "fmt": "xlsx"}
        )
        assert xlsx_resp.status_code == 200
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_resp.content))
        summary_text = "\n".join(
            str(cell.value) for row in wb["Summary"].iter_rows() for cell in row if cell.value
        )
        assert "POA&M / Risk Posture" in summary_text
        assert "Open POA&Ms" in summary_text
        assert "2" in summary_text  # open_poams value present somewhere on the sheet

        # docx: same section present.
        docx_resp = await c.get(
            "/api/reports/build", params={"organization_id": org_id, "fmt": "docx"}
        )
        assert docx_resp.status_code == 200
        doc = Document(io.BytesIO(docx_resp.content))
        docx_text = "\n".join(p.text for p in doc.paragraphs)
        assert "POA&M / Risk Posture" in docx_text
        assert "Open POA&Ms: 2" in docx_text
        assert "Overdue POA&Ms: 1" in docx_text
        assert "Risk-accepted POA&Ms: 1" in docx_text

        # csv: the risk preamble carries the same numbers ahead of the table.
        csv_resp = await c.get(
            "/api/reports/build", params={"organization_id": org_id, "fmt": "csv"}
        )
        assert csv_resp.status_code == 200
        csv_text = csv_resp.text
        assert "POA&M / Risk Posture (reconciled to dashboard)" in csv_text
        assert "Open POA&Ms,2" in csv_text
        assert "Overdue POA&Ms,1" in csv_text


@pytest.mark.asyncio
async def test_system_scoped_report_labels_risk_posture_as_organization_wide() -> None:
    """Finding: the posture/risk_summary block is always org-wide (org_summary),
    even when the report itself is scoped to one system via ``system_id`` — it
    must be labeled explicitly so a reader isn't misled into thinking it's
    limited to that system."""
    async with session_scope() as s:
        org = Organization(name="ReportScopeLabelOrg")
        s.add(org)
        await s.flush()
        sysm = System(organization_id=org.id, name="ReportScopeLabelSys", baseline="moderate")
        s.add(sysm)
        await s.flush()
        org_id, sys_id = org.id, sysm.id

    async with _client() as c:
        j = await c.get(
            "/api/reports/build",
            params={"organization_id": org_id, "system_id": sys_id, "fmt": "json"},
        )
        assert j.status_code == 200
        body = j.json()
        assert body["risk_summary"]["scope"] == "organization"

        csv_resp = await c.get(
            "/api/reports/build",
            params={"organization_id": org_id, "system_id": sys_id, "fmt": "csv"},
        )
        assert "Scope,Organization-wide" in csv_resp.text

        xlsx_resp = await c.get(
            "/api/reports/build",
            params={"organization_id": org_id, "system_id": sys_id, "fmt": "xlsx"},
        )
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_resp.content))
        summary_text = "\n".join(
            str(cell.value) for row in wb["Summary"].iter_rows() for cell in row if cell.value
        )
        assert "Organization-wide" in summary_text

        docx_resp = await c.get(
            "/api/reports/build",
            params={"organization_id": org_id, "system_id": sys_id, "fmt": "docx"},
        )
        doc = Document(io.BytesIO(docx_resp.content))
        docx_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Organization-wide" in docx_text


# --- CISO-10: AI/last-editor provenance flag on report rows ----------------


@pytest.mark.asyncio
async def test_report_flags_ai_sourced_rows() -> None:
    async with session_scope() as s:
        org = Organization(name="ReportAiOrg")
        s.add(org)
        await s.flush()
        sysm = System(organization_id=org.id, name="ReportAiSys", baseline="moderate")
        s.add(sysm)
        await s.flush()
        ctrl_draft = Control(identifier="AIPROV-01", control_name="AI-drafted control")
        ctrl_human = Control(identifier="AIPROV-02", control_name="Human-reviewed control")
        s.add_all([ctrl_draft, ctrl_human])
        await s.flush()
        s.add_all(
            [
                ControlImplementation(
                    system_id=sysm.id, control_id=ctrl_draft.id, status="implemented"
                ),
                ControlImplementation(
                    system_id=sysm.id, control_id=ctrl_human.id, status="implemented"
                ),
            ]
        )
        proj = SSPProject(organization_id=org.id, system_id=sysm.id, customer_name="AI Prov Co")
        s.add(proj)
        await s.flush()
        s.add_all(
            [
                SSPControlEntry(
                    project_id=proj.id,
                    control_id="AIPROV-01",
                    part_narratives=[
                        {"label": "Implementation", "text": DRAFT_PREFIX + "AI-drafted narrative."}
                    ],
                ),
                SSPControlEntry(
                    project_id=proj.id,
                    control_id="AIPROV-02",
                    part_narratives=[
                        {"label": "Implementation", "text": "Human-authored narrative."}
                    ],
                ),
            ]
        )
        org_id, sys_id = org.id, sysm.id

    async with _client() as c:
        j = await c.get(
            "/api/reports/build",
            params={"organization_id": org_id, "system_id": sys_id, "fmt": "json"},
        )
        assert j.status_code == 200
        rows = {row["identifier"]: row for row in j.json()["rows"]}
        assert rows["AIPROV-01"]["ai_sourced"] is True
        assert rows["AIPROV-02"]["ai_sourced"] is False

        xlsx_resp = await c.get(
            "/api/reports/build",
            params={"organization_id": org_id, "system_id": sys_id, "fmt": "xlsx"},
        )
        assert xlsx_resp.status_code == 200
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_resp.content))
        ws = wb["Controls"]
        headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        ai_col = headers.index("AI-Sourced")
        by_identifier = {row[0].value: row for row in ws.iter_rows(min_row=2)}
        assert by_identifier["AIPROV-01"][ai_col].value == "Yes"
        assert by_identifier["AIPROV-02"][ai_col].value in (None, "")
