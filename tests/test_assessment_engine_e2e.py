"""End-to-end: catalog objectives -> prepared evidence -> proposals -> acceptance -> SAR.

Every other assessment-engine test module verifies one stage against hand-built
fixtures for the *next* stage's input shape: ``test_assessment_objectives.py``
reads the catalog directly, ``test_assessment_evaluate.py`` seeds ``PrepUnit``
rows by hand, ``test_assessment_jobs.py`` and ``test_assessment_acceptance.py``
monkeypatch ``evaluate_objective`` itself rather than letting it run. That is
deliberate and correct for pinning each stage's own contract in isolation, but
none of it proves the stages actually chain end to end -- and a prior "end to
end" test in this codebase (slice 1's prep test) shipped exactly that gap: it
pre-marked four of five pipeline stages ``complete`` and hand-inserted the rows
they would have produced, so it looked like a full-chain test while only ever
exercising one stage. A reviewer caught it; this module exists so the same
mistake cannot happen again for the assessment engine.

This module seeds nothing but a real catalog control (with real sub-clause
objective rows, mirroring ``test_assessment_objectives.py``'s ``ZQ-`` fixture
shape but under its own ``ZAE-`` prefix -- no other module uses it) and a real
evidence document. It then drives the **real** slice-1 preparation pipeline
(``ccf.prep.pipeline.advance``, reusing ``test_prep_pipeline_e2e.py``'s
approach) to produce genuine ``prep_units``, opens a genuine
``AssessmentJob`` via ``enqueue_control``, drains it with ``run_once`` with
every proposal and job state left at its model default, accepts the resulting
proposal, and renders a real SAR docx from the accepted result -- asserting the
citation chain and the rendered output at every link, not just the final
state. The only two things monkeypatched anywhere in this module are the two
genuine external boundaries: the AI gateway's ``generate_structured`` and
``embed`` calls, exactly as ``test_prep_pipeline_e2e.py`` does for the prep
pipeline alone.
"""

from __future__ import annotations

import hashlib
import io
import re
from typing import Any

import pytest
from docx import Document
from sqlalchemy import delete, select, text

from ccf.ai import gateway
from ccf.ai.providers.base import EmbedResponse
from ccf.assessment.engine import jobs
from ccf.assessment.engine.evaluate import PURPOSE as EVALUATE_PURPOSE
from ccf.assessment.engine.service import accept_control_proposal
from ccf.assessment.sar import generate_sar_docx
from ccf.assessment.seed import result_to_dict, summarize_results
from ccf.config import get_settings
from ccf.db import session_scope
from ccf.evidence import storage
from ccf.models import Assessment, AssessmentControlResult, Control, Organization, System
from ccf.models_assessment_engine import (
    AssessmentControlProposal,
    AssessmentJob,
    AssessmentObjectiveProposal,
)
from ccf.models_evidence import EvidenceObject, EvidenceVersion
from ccf.models_prep import PrepUnit
from ccf.prep import pipeline
from ccf.prep.classify import ACTION_KEY as CLASSIFY_PURPOSE

pytestmark = pytest.mark.usefixtures("fresh_engine")

#: A prefix no other test module uses -- ``test_assessment_objectives.py`` and
#: its siblings own ``ZQ-``, ``test_assessment_jobs.py`` owns ``AEJ-``,
#: ``test_prep_pipeline_e2e.py`` owns the real ``IA-02``.
_SEQ = "ZAE-70"

#: Same weighting the real ETL applies after every workbook ingest -- copied
#: from ``test_prep_pipeline_e2e.py`` rather than imported, so this module
#: seeds and cleans up its own catalog rows independently.
_SEARCH_VECTOR_SQL = (
    "UPDATE ccf.controls SET search_vector = "
    "setweight(to_tsvector('english', coalesce(identifier,'')), 'A') || "
    "setweight(to_tsvector('english', coalesce(control_name,'')), 'A') || "
    "setweight(to_tsvector('english', coalesce(assessment_objective,'')), 'B') || "
    "setweight(to_tsvector('english', coalesce(description,'')), 'C') || "
    "setweight(to_tsvector('english', coalesce(discussion,'')), 'D')"
)


@pytest.fixture(autouse=True)
def _local_evidence_dir(tmp_path: Any, monkeypatch: pytest.MonkeyPatch) -> Any:
    monkeypatch.setenv("CCF_EVIDENCE_BACKEND", "local")
    monkeypatch.setenv("CCF_EVIDENCE_LOCAL_DIR", str(tmp_path / "evidence"))
    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


@pytest.fixture(autouse=True)
async def _catalog_rows() -> Any:
    """Seed one addressable control plus two real sub-clause objective rows.

    The parent row (``control_name`` populated) is what the prep pipeline's
    screen/classify stages can actually surface and cite -- mirrors
    ``test_prep_pipeline_e2e.py``'s ``IA-02`` fixture content, MFA phrasing
    and all, so the same lexical-overlap dynamics that module already proved
    clear the default screen threshold apply here too (its ``assessment_objective``
    holds the same MFA phrasing that fixture used, not a bare "Determine if:"
    header -- ``objectives_for`` never reads it for this row regardless,
    since it excludes every row with a non-null ``control_name``, so this is
    free to carry real signal instead of a header). The two sub-clause rows
    (``control_name`` NULL) are what ``objectives_for`` reads; ``score_line``
    explicitly excludes rows with no ``control_name``, so they cannot dilute
    screening.
    """
    async with session_scope() as s:
        await s.execute(delete(Control).where(Control.sequence_control == _SEQ))
        s.add(
            Control(
                identifier=_SEQ,
                sequence_control=_SEQ,
                control_name="Identification and Authentication (Organizational Users)",
                description=(
                    "Uniquely identify and authenticate organizational users and associate "
                    "that unique identification with processes acting on behalf of those "
                    "users. Multifactor authentication is required for network access to "
                    "privileged and non-privileged accounts."
                ),
                assessment_objective=(
                    "multifactor authentication is implemented for network access to "
                    "privileged accounts; multifactor authentication is implemented for "
                    "network access to non-privileged accounts"
                ),
                discussion=(
                    "Organizational users include employees or individuals considered to "
                    "have equivalent status. Authentication of user identities is "
                    "accomplished through passwords, tokens, or multifactor authentication "
                    "mechanisms such as personal identity verification cards."
                ),
                source_row=1,
            )
        )
        s.add(
            Control(
                identifier=f"{_SEQ}-ao1",
                sequence_control=_SEQ,
                assessment_objective=(
                    "multifactor authentication is implemented for network access to "
                    "privileged accounts;"
                ),
                source_row=2,
            )
        )
        s.add(
            Control(
                identifier=f"{_SEQ}-ao2",
                sequence_control=_SEQ,
                assessment_objective=(
                    "multifactor authentication is implemented for network access to "
                    "non-privileged accounts;"
                ),
                source_row=3,
            )
        )
        await s.flush()
        await s.execute(text(_SEARCH_VECTOR_SQL))
    try:
        yield
    finally:
        async with session_scope() as s:
            await s.execute(delete(Control).where(Control.sequence_control == _SEQ))


@pytest.fixture(autouse=True)
async def _isolate_job_queue() -> Any:
    """``assessment_jobs`` is a genuinely global queue with no organization
    filter (see ``test_assessment_jobs.py``'s identical fixture). Wiped
    unconditionally before and after so a stray pending row left by another
    module cannot be claimed alongside this test's own job and throw off the
    exact ``run_once`` stats asserted below.
    """

    async def _wipe() -> None:
        async with session_scope() as s:
            await s.execute(delete(AssessmentJob))

    await _wipe()
    yield
    await _wipe()


async def _org(name: str) -> int:
    async with session_scope() as s:
        org = Organization(name=name)
        s.add(org)
        await s.flush()
        return int(org.id)


async def _evidence_version(org_id: int, payload: bytes, filename: str) -> int:
    digest = hashlib.sha256(payload).hexdigest()
    ref = storage.get_backend().put(digest, payload, "text/plain")
    async with session_scope() as s:
        obj = EvidenceObject(organization_id=org_id, title=filename)
        s.add(obj)
        await s.flush()
        ver = EvidenceVersion(
            evidence_object_id=obj.id,
            version=1,
            sha256=digest,
            media_type="text/plain",
            size_bytes=len(payload),
            filename=filename,
            storage_backend="local",
            storage_ref=ref,
        )
        s.add(ver)
        await s.flush()
        return int(ver.id)


def _fake_embed(dim: int = 1024) -> Any:
    """Serve the prep pipeline's own embed stage; simulate an outage for retrieval.

    A constant vector for every text would make ``retriever._vector_ids``
    trivially "find" this test's single ``PrepUnit`` for *any* query -- with
    only one candidate row in the whole corpus, cosine distance against an
    identical embedding still returns that one row as the top match,
    regardless of whether it is actually relevant. That would let the
    objective-evaluation citation chain pass even if the classify stage's
    tagging were completely broken, which is precisely the vacuous-test shape
    this module exists to avoid (see the module docstring).

    So this only fabricates a usable vector for ``purpose="prep.embed"`` (the
    prep pipeline's own embed stage, which must still succeed for
    ``pipeline.advance`` to reach ``stage_embed == "complete"``). For
    ``purpose="prep.retrieve"`` (``ccf.prep.retriever``'s call, made once per
    objective evaluation) it raises the same ``GatewayError`` a real provider
    outage would -- ``retriever._vector_ids`` already degrades that to an
    empty vector ranking, by design (see that module's docstring). Combined
    with this fixture's evidence wording, which shares only some -- not
    all -- of each objective's lexemes with ``websearch_to_tsquery``'s
    AND semantics (so lexical retrieval also finds nothing), the only
    remaining path to this test's evidence is ``retriever._tagged_ids`` --
    i.e. the citations asserted below are proof classify's
    ``PrepClassification.control_identifiers`` genuinely reached retrieval,
    not an artifact of a single-candidate corpus.
    """

    async def _embed(session: Any, org_id: int, *, texts: list[str], **kw: Any) -> EmbedResponse:
        if kw.get("purpose") == "prep.retrieve":
            raise gateway.GatewayError("simulated embedding provider outage at retrieval time")
        return EmbedResponse(
            vectors=[[0.01] * dim for _ in texts],
            model="text-embedding-3-small",
            input_tokens=len(texts),
        )

    return _embed


async def _fake_generate_structured(
    session: Any,
    org_id: int,
    *,
    prompt: str,
    schema: dict[str, Any],
    purpose: str,
    system: str | None = None,
    **kw: Any,
) -> dict[str, Any]:
    """Stub both AI-gateway generation boundaries this chain crosses.

    Dispatches on ``purpose`` because two entirely different callers share
    this one monkeypatch target: the prep pipeline's classify stage
    (``purpose == CLASSIFY_PURPOSE``) and the assessment engine's own
    objective evaluation (``purpose == EVALUATE_PURPOSE``). Both echo back
    exactly what the real prompt offered rather than asserting on any value of
    their own authority, so a broken handoff in either direction (screen's
    candidates not reaching classify, or retrieval's passages not reaching
    evaluation) shows up as an empty response here, not a hand-picked one.
    """
    if purpose == CLASSIFY_PURPOSE:
        first_line = prompt.splitlines()[0]
        offered = first_line.removeprefix("Candidate controls:").strip()
        candidates = (
            []
            if offered == "(none surfaced by screening)"
            else [c.strip() for c in offered.split(",")]
        )
        return {
            "control_identifiers": candidates,
            "artifact_type": "policy",
            "evidence_strength": "strong",
            "confidence": 0.9,
        }
    if purpose == EVALUATE_PURPOSE:
        # evaluate.build_prompt renders each offered passage as "[<unit_id>]
        # (page ...)" -- parsed back out here rather than hardcoded, so the
        # citations this test later asserts on are proof retrieval's output
        # actually reached the model, not an assertion this fake grants
        # itself.
        offered_ids = [int(m) for m in re.findall(r"\[(\d+)\]", prompt)]
        return {
            "verdict": "satisfied",
            "cited_unit_ids": offered_ids,
            "gaps": [],
            "contradictions": [],
            "rationale": "The cited passage demonstrates multifactor authentication is used.",
            "confidence": 0.85,
        }
    raise AssertionError(f"unexpected generate_structured purpose: {purpose!r}")


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


async def test_a_real_assessment_produces_accepted_findings_with_citations(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Catalog objectives -> prepared evidence -> proposals -> acceptance -> SAR."""
    org_id = await _org("assessment-engine-e2e")
    payload = (
        b"Administrators must use multifactor authentication for network access.\n"
        b"The quick brown fox jumped over the lazy dog.\n"
    )
    version_id = await _evidence_version(org_id, payload, "access-control.txt")

    monkeypatch.setattr(gateway, "generate_structured", _fake_generate_structured)
    monkeypatch.setattr(gateway, "embed", _fake_embed())

    # --- Stage 1: the real slice-1 preparation pipeline, start to finish. ---
    async with session_scope() as s:
        run = await pipeline.create_run(
            s, organization_id=org_id, source_kind="evidence_version", source_id=version_id
        )
        # Nothing pre-marked -- every stage starts at its model default.
        assert run.stage_parse == "pending"
        assert run.stage_screen == "pending"
        assert run.stage_expand == "pending"
        assert run.stage_classify == "pending"
        assert run.stage_embed == "pending"

        await pipeline.advance(s, run)

        assert run.status == "complete", run.error
        assert run.stage_parse == "complete"
        assert run.stage_screen == "complete"
        assert run.stage_expand == "complete"
        assert run.stage_classify == "complete"
        assert run.stage_embed == "complete"

    async with session_scope() as s:
        units = (
            (await s.execute(select(PrepUnit).where(PrepUnit.organization_id == org_id)))
            .scalars()
            .all()
        )
    assert units, "the pipeline must have produced at least one genuine prep_units row"
    unit_ids = {int(u.id) for u in units}

    # --- Stage 2: a real Assessment over the same organization's system. ---
    async with session_scope() as s:
        system = System(organization_id=org_id, name="assessment-engine-e2e-system")
        s.add(system)
        await s.flush()
        assessment = Assessment(
            system_id=system.id, name="assessment-engine-e2e-assessment", kind="self"
        )
        s.add(assessment)
        await s.flush()
        assessment_id = int(assessment.id)

    # --- Stage 3: enqueue + drain -- nothing pre-marked here either. ---
    async with session_scope() as s:
        job = await jobs.enqueue_control(s, assessment_id=assessment_id, control_identifier=_SEQ)
        assert job.status == "pending"

    async with session_scope() as s:
        stats = await jobs.run_once(s, worker="e2e", limit=5)
    assert stats == {"claimed": 1, "finished": 1, "failed": 0}, stats

    # --- Stage 4: assert the whole citation chain, not just the endpoints. ---
    async with session_scope() as s:
        proposal = (
            await s.execute(
                select(AssessmentControlProposal).where(
                    AssessmentControlProposal.assessment_id == assessment_id
                )
            )
        ).scalar_one()
        assert proposal.control_identifier == _SEQ
        assert proposal.state == "complete"
        assert proposal.proposed_finding == "satisfied"
        proposal_id = int(proposal.id)

        objectives = (
            (
                await s.execute(
                    select(AssessmentObjectiveProposal)
                    .where(AssessmentObjectiveProposal.control_proposal_id == proposal_id)
                    .order_by(AssessmentObjectiveProposal.sort_order)
                )
            )
            .scalars()
            .all()
        )

    assert [o.label for o in objectives] == [f"{_SEQ}a", f"{_SEQ}b"], (
        "exactly one AssessmentObjectiveProposal per catalog objective, in sort_order"
    )
    assert [o.sort_order for o in objectives] == [0, 1]
    assert [o.verdict for o in objectives] == ["satisfied", "satisfied"]

    cited_ids: set[int] = set()
    for o in objectives:
        assert o.cited_unit_ids, f"objective {o.label} carries no citations at all"
        for uid in o.cited_unit_ids:
            assert uid in unit_ids, (
                f"objective {o.label} cites unit {uid}, which is not a real prep_units row "
                "produced by this run"
            )
            cited_ids.add(int(uid))

    async with session_scope() as s:
        cited_units = (
            (await s.execute(select(PrepUnit).where(PrepUnit.id.in_(cited_ids)))).scalars().all()
        )
    assert len(cited_units) == len(cited_ids), "every cited_unit_id must resolve to a live row"
    for u in cited_units:
        assert u.page_numbers, f"cited unit {u.id} has no page_numbers -- not auditable"

    # --- Stage 5: acceptance -- the human gate into AssessmentControlResult. ---
    async with session_scope() as s:
        result = await accept_control_proposal(s, proposal_id, accepted_by="e2e@example.com")
        assert len(result.objective_findings) == 2
        assert result.finding == "satisfied"
        result_dict = result_to_dict(result)
        summary = summarize_results([result])

    # Re-fetch from a fresh session: proves the accepted result is durable,
    # not merely visible within the mutating session above.
    async with session_scope() as s:
        stored = (
            await s.execute(
                select(AssessmentControlResult).where(
                    AssessmentControlResult.assessment_id == assessment_id,
                    AssessmentControlResult.control_id == _SEQ,
                )
            )
        ).scalar_one()
        assert len(stored.objective_findings) == 2
        assert stored.finding == "satisfied"

    # --- Stage 6: render the real SAR and check both objective labels land. ---
    meta = {
        "customer_name": "Assessment Engine E2E Org",
        "system_name": "assessment-engine-e2e-system",
        "assessment_name": "assessment-engine-e2e-assessment",
        "kind": "self",
        "assessor": "e2e@example.com",
        "period": "2026-01-01 to ongoing",
        "date": "08/11/2026",
    }
    data = generate_sar_docx(meta, summary, [result_dict])
    text_body = _docx_text(data)

    assert f"{_SEQ}a" in text_body
    assert f"{_SEQ}b" in text_body
