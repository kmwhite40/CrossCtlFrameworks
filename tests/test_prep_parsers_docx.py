"""DOCX parser — heading breadcrumbs and table cell headers."""

from __future__ import annotations

import io
import logging

from docx import Document

from ccf.prep.parsers import dispatch
from ccf.prep.parsers.docx import parse_docx


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Access Control", level=1)
    doc.add_heading("Account Management", level=2)
    doc.add_paragraph("Accounts are reviewed by the system owner.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Activity"
    table.cell(0, 1).text = "Review Frequency"
    table.cell(1, 0).text = "Privileged account review"
    table.cell(1, 1).text = "Quarterly"
    doc.add_heading("Audit and Accountability", level=1)
    doc.add_paragraph("Audit logs are retained for one year.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_paragraph_inherits_its_heading_breadcrumb() -> None:
    doc = parse_docx(_docx_bytes(), "policy.docx")
    line = next(x for x in doc.iter_lines() if x.content.startswith("Accounts are reviewed"))
    assert line.section_path == "Access Control > Account Management"


def test_breadcrumb_resets_when_a_higher_level_heading_appears() -> None:
    """A level-1 heading must clear the stale level-2 crumb, not append to it."""
    doc = parse_docx(_docx_bytes(), "policy.docx")
    line = next(x for x in doc.iter_lines() if x.content.startswith("Audit logs"))
    assert line.section_path == "Audit and Accountability"


def test_table_cell_carries_its_column_header_and_coordinates() -> None:
    doc = parse_docx(_docx_bytes(), "policy.docx")
    line = next(x for x in doc.iter_lines() if x.content == "Quarterly")
    assert line.cell_label == "Review Frequency"
    assert line.block_type == "table_cell"
    assert (line.row_index, line.col_index) == (1, 1)
    assert line.table_id is not None


def test_header_row_cells_have_no_self_referential_label() -> None:
    doc = parse_docx(_docx_bytes(), "policy.docx")
    line = next(x for x in doc.iter_lines() if x.content == "Review Frequency")
    assert line.cell_label is None


def test_dispatch_routes_docx() -> None:
    doc = dispatch(_docx_bytes(), "policy.docx", None)
    assert doc.parser_name == "docx"
    assert doc.success


def test_corrupt_docx_returns_an_error_document_rather_than_raising() -> None:
    doc = parse_docx(b"not a real docx", "broken.docx")
    assert not doc.success
    assert doc.error is not None


def test_corrupt_docx_failure_path_survives_an_enabled_logger() -> None:
    """Regression for the reserved-LogRecord-attribute defect.

    tests/conftest.py runs Alembic migrations, and migrations/env.py calls
    logging.config.fileConfig(..., disable_existing_loggers=True), which disables
    this module's stdlib logger for the whole pytest session. A logger.warning()
    call on a *disabled* logger short-circuits before building a LogRecord, so a
    bug in the log call itself (e.g. passing `extra={"filename": ...}`, which
    collides with LogRecord's own reserved `filename` attribute and raises
    KeyError from stdlib logging) would never fire and this suite would pass
    while the same call raises outside pytest. Force the logger enabled so this
    test genuinely exercises the log call instead of being masked by that side
    effect.
    """
    logger = logging.getLogger("ccf.prep.parsers.docx")
    original = logger.disabled
    logger.disabled = False
    try:
        doc = parse_docx(b"not a real docx", "broken.docx")
    finally:
        logger.disabled = original
    assert not doc.success
    assert doc.error is not None
