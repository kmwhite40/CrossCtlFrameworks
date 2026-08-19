"""FedRAMP 20x machine-readable authorization-package foundation.

Assembles the CSO profile, KSI catalog + latest validation results, evidence
references, dependency inventory, assessor review status, exceptions, POA&M
linkage, a continuous-monitoring summary, and the system authorization status
into a single structured document. Renders to JSON, Markdown, and an
OSCAL-*shaped* JSON structure.

NOTE: the OSCAL-ready output follows OSCAL naming/nesting conventions. It is
validated against a Concord-maintained OSCAL-*subset* JSON Schema
(``oscal_ksi.schema.json``) when ``jsonschema`` is installed — this is NOT the
official OSCAL schema and does not claim full OSCAL compliance. FedRAMP 20x package
requirements are still evolving; treat this as a foundation, not a submission artifact.
"""

from __future__ import annotations

import io
import json
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from ..constants import POAM_ACTIVE_STATUSES
from ..models import (
    KSI,
    POAM,
    FedRAMP20xProfile,
    FedRAMPDependency,
    KSIAssessorReview,
    KSIException,
    KSIState,
    KSIValidationResult,
    System,
)
from .readiness import score_system


async def _latest_validation(
    session: AsyncSession, system_id: int
) -> dict[int, KSIValidationResult]:
    rows = (
        (
            await session.execute(
                select(KSIValidationResult)
                .where(KSIValidationResult.system_id == system_id)
                .order_by(KSIValidationResult.validated_at.desc())
            )
        )
        .scalars()
        .all()
    )
    latest: dict[int, KSIValidationResult] = {}
    for r in rows:
        latest.setdefault(r.ksi_id, r)  # first seen == newest (desc order)
    return latest


async def build_package(session: AsyncSession, *, system_id: int) -> dict[str, Any]:
    """Build the structured 20x authorization package for a system."""
    system = (
        await session.execute(select(System).where(System.id == system_id))
    ).scalar_one_or_none()
    if system is None:
        raise ValueError(f"system {system_id} not found")

    profile = (
        await session.execute(
            select(FedRAMP20xProfile).where(FedRAMP20xProfile.system_id == system_id)
        )
    ).scalar_one_or_none()
    ksis = list((await session.execute(select(KSI).order_by(KSI.sort_order))).scalars().all())
    states = {
        s.ksi_id: s
        for s in (await session.execute(select(KSIState).where(KSIState.system_id == system_id)))
        .scalars()
        .all()
    }
    latest_val = await _latest_validation(session, system_id)
    reviews_rows = (
        (
            await session.execute(
                select(KSIAssessorReview).where(KSIAssessorReview.system_id == system_id)
            )
        )
        .scalars()
        .all()
    )
    reviews: dict[int, KSIAssessorReview] = {}
    for rv in reviews_rows:
        prev = reviews.get(rv.ksi_id)
        if prev is None or rv.reviewed_at >= prev.reviewed_at:
            reviews[rv.ksi_id] = rv
    deps = list(
        (
            await session.execute(
                select(FedRAMPDependency).where(FedRAMPDependency.system_id == system_id)
            )
        )
        .scalars()
        .all()
    )
    exceptions = list(
        (await session.execute(select(KSIException).where(KSIException.system_id == system_id)))
        .scalars()
        .all()
    )
    poams = list(
        (
            await session.execute(
                select(POAM)
                .options(selectinload(POAM.milestones))
                .where(POAM.system_id == system_id)
            )
        )
        .scalars()
        .all()
    )
    # Readiness without persisting a snapshot (this is a read/export path).
    readiness = await score_system(session, system_id=system_id, persist=False)

    ksi_entries: list[dict[str, Any]] = []
    conmon_counts: dict[str, int] = {}
    for k in ksis:
        st = states.get(k.id)
        val = latest_val.get(k.id)
        review = reviews.get(k.id)
        status = st.status if st else "not_tested"
        conmon_counts[status] = conmon_counts.get(status, 0) + 1
        ksi_entries.append(
            {
                "identifier": k.identifier,
                "name": k.name,
                "category": k.category,
                "category_name": k.category_name,
                "expected_outcome": k.expected_outcome,
                "validation_method": k.validation_method,
                "validation_frequency": k.validation_frequency,
                "automation_level": k.automation_level,
                "nist_refs": list(k.nist_refs or []),
                "cmmc_refs": list(k.cmmc_refs or []),
                "status": status,
                "last_validated_at": st.last_validated_at.isoformat()
                if st and st.last_validated_at
                else None,
                "next_validation_due": st.next_validation_due.isoformat()
                if st and st.next_validation_due
                else None,
                "validation": None
                if val is None
                else {
                    "status": val.status,
                    "confidence": val.confidence,
                    "source": val.source,
                    "evidence_refs": list(val.evidence_refs or []),
                    "failure_reason": val.failure_reason,
                    "remediation_hint": val.remediation_hint,
                    "assessor_review_required": val.assessor_review_required,
                },
                "assessor_review": None
                if review is None
                else {
                    "status": review.status,
                    "assessor": review.assessor,
                    "finding": review.finding,
                },
            }
        )

    return {
        "generated_at": datetime.now(UTC).isoformat(),
        "disclaimer": (
            "Foundation output. Not an official FedRAMP submission. OSCAL-shaped "
            "sections are validated against a Concord OSCAL-subset schema, not the "
            "official OSCAL schema. KSI catalog is representative and evolving."
        ),
        "system": {
            "id": system.id,
            "name": system.name,
            "description": system.description,
            "baseline": system.baseline,
            "ato_status": system.ato_status,
            "ato_expires_on": system.ato_expires_on.isoformat() if system.ato_expires_on else None,
        },
        "cloud_service_offering": None
        if profile is None
        else {
            "service_name": profile.service_name,
            "service_description": profile.service_description,
            "deployment_model": profile.deployment_model,
            "cloud_environment": profile.cloud_environment,
            "boundary_description": profile.boundary_description,
            "data_types": list(profile.data_types or []),
            "federal_data": profile.federal_data,
            "cui_support": profile.cui_support,
            "external_services": list(profile.external_services or []),
            "responsibilities": {
                "customer": profile.customer_responsibilities,
                "provider": profile.provider_responsibilities,
                "shared": profile.shared_responsibilities,
            },
            "boundaries": {
                "cryptographic": profile.cryptographic_boundary,
                "logging": profile.logging_boundary,
                "incident_response": profile.incident_response_boundary,
                "backup_recovery": profile.backup_recovery_boundary,
                "administrative_access": profile.administrative_access_boundary,
            },
            "readiness_status": profile.readiness_status,
        },
        "readiness": readiness,
        "ksis": ksi_entries,
        "dependencies": [
            {
                "name": d.name,
                "provider": d.provider,
                "service_type": d.service_type,
                "fedramp_status": d.fedramp_status,
                "authorization_level": d.authorization_level,
                "impact_level": d.impact_level,
                "boundary_role": d.boundary_role,
                "dependency_risk": d.dependency_risk,
                "inherited_controls": list(d.inherited_controls or []),
            }
            for d in deps
        ],
        "exceptions": [
            {
                "ksi_id": e.ksi_id,
                "rationale": e.rationale,
                "status": e.status,
                "expires_on": e.expires_on.isoformat() if e.expires_on else None,
            }
            for e in exceptions
        ],
        "poams": [
            {
                "id": p.id,
                "title": p.title,
                "severity": p.severity,
                "status": p.status,
                "due_on": p.due_on.isoformat() if p.due_on else None,
            }
            for p in poams
        ],
        "continuous_monitoring": {
            "ksi_state_counts": conmon_counts,
            "validation_history_count": len(latest_val),
        },
    }


def render_markdown(pkg: dict[str, Any]) -> str:
    """Render the package as a human-readable Markdown brief."""
    sys = pkg["system"]
    r = pkg["readiness"]
    cso = pkg.get("cloud_service_offering")

    def pct(key: str) -> str:
        v = r.get(key)
        return f"{v}%" if v is not None else "n/a"

    lines: list[str] = [
        f"# FedRAMP 20x Package — {sys['name']}",
        "",
        f"_Generated {pkg['generated_at']}_",
        "",
        f"> {pkg['disclaimer']}",
        "",
        "## System",
        f"- **Baseline:** {sys.get('baseline') or 'n/a'}",
        f"- **ATO status:** {sys.get('ato_status') or 'n/a'}",
        "",
        "## 20x Readiness",
        f"- **Overall:** {r['readiness_pct']}%  ·  **Status:** {r['status']}",
        f"- **KSI pass rate:** {pct('ksi_pass_rate')}  ·  "
        f"**Automation coverage:** {pct('automation_coverage')}",
        f"- **Evidence completeness:** {pct('evidence_completeness')}  ·  "
        f"**ConMon coverage:** {pct('conmon_coverage')}",
        f"- **Assessor completion:** {pct('assessor_completion')}  ·  "
        f"**Dependency readiness:** {pct('dependency_readiness')}",
        f"- **Open exceptions:** {r['open_exceptions']}  ·  "
        f"**High-risk findings:** {r['high_risk_findings']}  ·  "
        f"**Manual-review burden:** {r['manual_review_burden']}",
        "",
    ]
    if cso:
        lines += [
            "## Cloud Service Offering",
            f"- **Service:** {cso.get('service_name') or 'n/a'}",
            f"- **Deployment / environment:** {cso.get('deployment_model') or 'n/a'} / "
            f"{cso.get('cloud_environment') or 'n/a'}",
            "",
        ]
    lines += [
        "## Key Security Indicators",
        "",
        "| KSI | Category | Status | Assessor |",
        "|---|---|---|---|",
    ]
    for k in pkg["ksis"]:
        rv = (k.get("assessor_review") or {}).get("status", "not_reviewed")
        lines.append(f"| {k['identifier']} | {k['category']} | {k['status']} | {rv} |")
    lines += ["", "## Dependencies", ""]
    if pkg["dependencies"]:
        lines += ["| Name | Provider | FedRAMP status | Risk |", "|---|---|---|---|"]
        for d in pkg["dependencies"]:
            lines.append(
                f"| {d['name']} | {d.get('provider') or ''} | {d['fedramp_status']} | "
                f"{d.get('dependency_risk') or ''} |"
            )
    else:
        lines.append("_No dependencies recorded._")
    open_poams = sum(1 for p in pkg["poams"] if p["status"] in POAM_ACTIVE_STATUSES)
    lines += ["", f"**Open POA&Ms:** {open_poams}", ""]
    return "\n".join(lines)


def to_oscal_shaped(pkg: dict[str, Any]) -> dict[str, Any]:
    """Emit an OSCAL-*shaped* structure (assessment-results flavored).

    Not schema-validated; naming mirrors OSCAL to ease a future real export.
    """
    findings = []
    observations = []
    for k in pkg["ksis"]:
        observations.append(
            {
                "uuid": f"obs-{k['identifier']}",
                "title": f"{k['identifier']} — {k['name']}",
                "props": [
                    {"name": "ksi-category", "value": k["category"]},
                    {"name": "validation-status", "value": k["status"]},
                    {"name": "automation-level", "value": k.get("automation_level") or "manual"},
                ],
                "relevant-evidence": [
                    {"description": ref}
                    for ref in ((k.get("validation") or {}).get("evidence_refs") or [])
                ],
                "related-controls": list(k.get("nist_refs") or []),
            }
        )
        if k["status"] in ("fail", "warn"):
            findings.append(
                {
                    "uuid": f"finding-{k['identifier']}",
                    "title": f"{k['identifier']} not satisfied",
                    "target": {"type": "ksi", "target-id": k["identifier"], "status": k["status"]},
                    "description": (k.get("validation") or {}).get("failure_reason") or "",
                }
            )
    return {
        "assessment-results": {
            "uuid": f"cso-{pkg['system']['id']}",
            "metadata": {
                "title": f"FedRAMP 20x KSI Assessment — {pkg['system']['name']}",
                "last-modified": pkg["generated_at"],
                "oscal-version": "1.1.2",
                "remarks": pkg["disclaimer"],
            },
            "results": [
                {
                    "uuid": "ksi-result-1",
                    "title": "KSI Validation Results",
                    "props": [
                        {"name": "readiness-pct", "value": str(pkg["readiness"]["readiness_pct"])},
                        {"name": "readiness-status", "value": pkg["readiness"]["status"]},
                    ],
                    "observations": observations,
                    "findings": findings,
                }
            ],
        }
    }


SCHEMA_PATH = Path(__file__).with_name("oscal_ksi.schema.json")


def validation_mode() -> str:
    """Which validation backend ``validate_oscal`` will use: 'jsonschema' or 'structural'."""
    try:
        import jsonschema  # noqa: F401,PLC0415

        return "jsonschema" if SCHEMA_PATH.is_file() else "structural"
    except Exception:
        return "structural"


def _validate_oscal_schema(doc: dict[str, Any]) -> list[str]:
    """Validate against the bundled JSON Schema (Concord subset, not official OSCAL)."""
    import jsonschema  # noqa: PLC0415

    schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
    validator = jsonschema.Draft202012Validator(schema)
    return [
        f"{'/'.join(str(p) for p in e.absolute_path) or 'root'}: {e.message}"
        for e in sorted(validator.iter_errors(doc), key=lambda e: list(e.absolute_path))
    ]


def validate_oscal(doc: dict[str, Any]) -> list[str]:
    """Validate an OSCAL-shaped assessment-results document.

    Prefers JSON Schema validation against the bundled Concord subset schema
    (``oscal_ksi.schema.json``) when ``jsonschema`` is installed; otherwise falls
    back to built-in structural checks. Either way this validates the subset
    Concord emits — it is NOT a claim of full official OSCAL conformance. Returns a
    list of human-readable error strings; empty means valid.
    """
    if validation_mode() == "jsonschema":
        return _validate_oscal_schema(doc)
    return _validate_oscal_structural(doc)


def _validate_oscal_structural(doc: dict[str, Any]) -> list[str]:
    """Fallback structural validation used when ``jsonschema`` is unavailable."""
    errors: list[str] = []

    def req(obj: dict[str, Any], key: str, where: str, typ: type | tuple[type, ...]) -> Any:
        if key not in obj:
            errors.append(f"{where}: missing required key '{key}'")
            return None
        if not isinstance(obj[key], typ):
            errors.append(f"{where}.{key}: expected {typ}, got {type(obj[key]).__name__}")
            return None
        return obj[key]

    if not isinstance(doc, dict):
        return ["root: document must be an object"]
    ar = req(doc, "assessment-results", "root", dict)
    if ar is None:
        return errors
    req(ar, "uuid", "assessment-results", str)
    meta = req(ar, "metadata", "assessment-results", dict)
    if isinstance(meta, dict):
        for key in ("title", "last-modified", "oscal-version"):
            req(meta, key, "assessment-results.metadata", str)
    results = req(ar, "results", "assessment-results", list)
    if isinstance(results, list):
        if not results:
            errors.append("assessment-results.results: must contain at least one result")
        for i, res in enumerate(results):
            where = f"results[{i}]"
            if not isinstance(res, dict):
                errors.append(f"{where}: must be an object")
                continue
            req(res, "uuid", where, str)
            req(res, "title", where, str)
            for coll in ("observations", "findings"):
                items = res.get(coll, [])
                if not isinstance(items, list):
                    errors.append(f"{where}.{coll}: expected list")
                    continue
                for j, it in enumerate(items):
                    if not isinstance(it, dict) or "uuid" not in it:
                        errors.append(f"{where}.{coll}[{j}]: missing 'uuid'")
    return errors


def to_docx(pkg: dict[str, Any]) -> bytes:
    """Render the package as a Word (.docx) brief. Reuses python-docx (a core dep)."""
    from docx import Document  # noqa: PLC0415 — heavy optional import kept local
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    doc.add_heading(f"FedRAMP 20x Package — {pkg['system']['name']}", level=0)
    doc.add_paragraph(f"Generated {pkg['generated_at']}").runs[0].italic = True
    doc.add_paragraph(pkg["disclaimer"])

    r = pkg["readiness"]
    doc.add_heading("20x Readiness", level=1)
    for label, key in (
        ("Overall readiness", "readiness_pct"),
        ("Lifecycle status", "status"),
        ("KSI pass rate", "ksi_pass_rate"),
        ("Automation coverage", "automation_coverage"),
        ("Evidence completeness", "evidence_completeness"),
        ("Assessor completion", "assessor_completion"),
        ("Dependency readiness", "dependency_readiness"),
        ("High-risk findings", "high_risk_findings"),
        ("Manual-review burden", "manual_review_burden"),
    ):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(str(r.get(key)))

    doc.add_heading("Key Security Indicators", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "KSI", "Category", "Status", "Assessor"
    for k in pkg["ksis"]:
        row = table.add_row().cells
        row[0].text = k["identifier"]
        row[1].text = k["category"]
        row[2].text = k["status"]
        row[3].text = (k.get("assessor_review") or {}).get("status", "not_reviewed")

    if pkg["dependencies"]:
        doc.add_heading("Authorized dependencies", level=1)
        for d in pkg["dependencies"]:
            doc.add_paragraph(
                f"{d['name']} ({d.get('provider') or 'n/a'}) — {d['fedramp_status']}",
                style="List Bullet",
            )

    for s in doc.styles:
        if s.name == "Normal":
            s.font.size = Pt(10)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def evidence_manifest(pkg: dict[str, Any]) -> list[dict[str, Any]]:
    """Flat manifest of every evidence reference cited across the package's KSIs."""
    manifest: list[dict[str, Any]] = []
    for k in pkg["ksis"]:
        for ref in ((k.get("validation") or {}).get("evidence_refs") or []):
            manifest.append({"ksi": k["identifier"], "evidence_ref": ref})
    return manifest


def to_bundle(pkg: dict[str, Any]) -> bytes:
    """Zip the full package: JSON, Markdown, OSCAL-shaped JSON, and an evidence manifest."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("package.json", json.dumps(pkg, indent=2, default=str))
        z.writestr("package.md", render_markdown(pkg))
        z.writestr("oscal.json", json.dumps(to_oscal_shaped(pkg), indent=2, default=str))
        z.writestr(
            "evidence-manifest.json", json.dumps(evidence_manifest(pkg), indent=2, default=str)
        )
        z.writestr("README.txt", pkg["disclaimer"] + "\n")
    return buf.getvalue()
