"""Word document parser.

Two structures matter and both are easy to lose. A paragraph's meaning depends on
the heading it sits under, so headings are tracked as a stack that truncates when
a higher level appears — otherwise "Accounts are reviewed quarterly" under
Audit inherits a stale Access Control crumb. And a table cell's meaning depends
on its column header: "Quarterly" is not evidence, "Review Frequency: Quarterly"
is.

python-docx exposes paragraphs and tables as separate collections, losing their
relative order. Document order is recovered by walking the body's XML children
and mapping each element back to its wrapper.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from ...logging import get_logger
from .base import ParsedBlock, ParsedCell, ParsedDocument, ParsedPage

log = get_logger(__name__)

PARSER_NAME = "docx"
MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _heading_level(paragraph: Paragraph) -> int | None:
    """Return the outline level of a heading paragraph, or ``None`` if it is body text."""
    style = (paragraph.style.name or "") if paragraph.style is not None else ""
    if not style.startswith("Heading"):
        return None
    tail = style.removeprefix("Heading").strip()
    return int(tail) if tail.isdigit() else 1


def _iter_body(document: DocxDocument) -> list[Paragraph | Table]:
    """Yield paragraphs and tables in true document order."""
    paragraphs = {p._element: p for p in document.paragraphs}
    tables = {t._element: t for t in document.tables}
    ordered: list[Paragraph | Table] = []
    for element in document.element.body.iterchildren():
        found = paragraphs.get(element) or tables.get(element)
        if found is not None:
            ordered.append(found)
    return ordered


def _table_blocks(table: Table, table_index: int, heading_path: list[str]) -> list[ParsedBlock]:
    """Flatten a table to one block per non-empty cell, carrying its column header."""
    table_id = f"t{table_index}"
    rows = table.rows
    if not rows:
        return []
    headers = [cell.text.strip() for cell in rows[0].cells]
    blocks: list[ParsedBlock] = []
    for row_index, row in enumerate(rows):
        cells = [
            ParsedCell(
                text=cell.text.strip(),
                row_index=row_index,
                col_index=col_index,
                header=headers[col_index] if col_index < len(headers) else None,
            )
            for col_index, cell in enumerate(row.cells)
        ]
        for cell in cells:
            if not cell.text:
                continue
            blocks.append(
                ParsedBlock(
                    block_id=f"{table_id}r{cell.row_index}c{cell.col_index}",
                    block_type="table_cell",
                    text=cell.text,
                    heading_path=list(heading_path),
                    row_index=cell.row_index,
                    col_index=cell.col_index,
                    table_id=table_id,
                    # The header row labels itself otherwise, which is noise.
                    cell_label=cell.header if row_index > 0 else None,
                    cells=cells,
                )
            )
    return blocks


def parse_docx(data: bytes, filename: str, media_type: str = MEDIA_TYPE) -> ParsedDocument:
    """Parse a .docx into heading-aware paragraph and table-cell blocks."""
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # any malformed file must not kill the run
        log.warning("prep.parse.docx_failed", filename=filename, error=str(exc))
        return ParsedDocument(
            filename=filename, media_type=media_type, parser_name=PARSER_NAME,
            error=f"could not open document: {exc}",
        )

    blocks: list[ParsedBlock] = []
    heading_path: list[str] = []
    table_index = 0
    paragraph_index = 0

    for item in _iter_body(document):
        if isinstance(item, Table):
            table_index += 1
            blocks.extend(_table_blocks(item, table_index, heading_path))
            continue

        text = item.text.strip()
        if not text:
            continue
        level = _heading_level(item)
        if level is not None:
            # Truncate to the parent depth before pushing, so a level-1 heading
            # discards any deeper crumbs left over from the previous section.
            del heading_path[level - 1 :]
            heading_path.append(text)
            block_type = "heading"
        else:
            block_type = "paragraph"
        paragraph_index += 1
        blocks.append(
            ParsedBlock(
                block_id=f"p{paragraph_index}",
                block_type=block_type,
                text=text,
                heading_path=list(heading_path),
            )
        )

    metadata: dict[str, Any] = {"block_count": len(blocks)}
    return ParsedDocument(
        filename=filename,
        media_type=media_type,
        parser_name=PARSER_NAME,
        pages=[ParsedPage(page_number=1, blocks=blocks, metadata=metadata)],
    )
