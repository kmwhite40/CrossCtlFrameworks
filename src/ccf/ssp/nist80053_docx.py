"""Render a NIST SP 800-53 Rev 5 SSP project to a FedRAMP-style ``.docx``.

Parallel to :func:`ccf.ssp.generator.generate_ssp_docx` (the CMMC Level 2 /
NIST SP 800-171 renderer) but for the 800-53r5 pipeline: entries come from
:func:`ccf.ssp.nist80053.build_80053_entries` / :func:`ccf.ssp.seed.seed_80053_project`
and are keyed by the real NIST 800-53 control id and family (``domain``), not
CMMC practice ids. Reuses the low-level table/shading helpers from
``generator.py`` rather than duplicating them.
"""

from __future__ import annotations

import io
from collections.abc import Mapping, Sequence
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..catalog.canonical import canonicalize
from .generator import ACCENT, HEADER_FILL, LABEL_FILL, _runs, _shade

# NIST SP 800-53 Rev 5 control family names, keyed by the two-letter family
# code that ``build_80053_entries`` stores as each entry's ``domain``.
FAMILY_NAMES: dict[str, str] = {
    "AC": "Access Control",
    "AT": "Awareness and Training",
    "AU": "Audit and Accountability",
    "CA": "Assessment, Authorization, and Monitoring",
    "CM": "Configuration Management",
    "CP": "Contingency Planning",
    "IA": "Identification and Authentication",
    "IR": "Incident Response",
    "MA": "Maintenance",
    "MP": "Media Protection",
    "PE": "Physical and Environmental Protection",
    "PL": "Planning",
    "PM": "Program Management",
    "PS": "Personnel Security",
    "PT": "PII Processing and Transparency",
    "RA": "Risk Assessment",
    "SA": "System and Services Acquisition",
    "SC": "System and Communications Protection",
    "SI": "System and Information Integrity",
    "SR": "Supply Chain Risk Management",
}

FAMILY_ORDER = sorted(FAMILY_NAMES)

_BASELINE_LABELS = {"low": "Low", "moderate": "Moderate", "high": "High"}


def _family_name(code: str) -> str:
    return FAMILY_NAMES.get(code, code)


def _baseline_label(project: Mapping[str, Any]) -> str:
    raw = str(project.get("baseline") or "").strip().lower()
    return _BASELINE_LABELS.get(raw, raw.title() if raw else "Not set")


def _sort_key(entry: Mapping[str, Any]) -> tuple[str, int, tuple[int, ...]]:
    cid = str(entry.get("control_id") or "")
    parsed = canonicalize(cid)
    if parsed is None:
        return (cid, 0, ())
    return (parsed.family, parsed.number, parsed.enhancements)


def _kv_table(doc: Any, rows: Sequence[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lc = table.cell(i, 0)
        _shade(lc, LABEL_FILL)
        _runs(lc, label, bold=True)
        _runs(table.cell(i, 1), str(value) if value else "—")


def _cover(doc: Any, project: Mapping[str, Any]) -> None:
    def centered(text: str, size: int, bold: bool, color: RGBColor | None = None) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color

    for _ in range(3):
        doc.add_paragraph()
    centered(project.get("title") or "System Security Plan", 26, True, ACCENT)
    centered("NIST SP 800-53 Rev 5", 15, True)
    baseline = _baseline_label(project)
    centered(
        f"FIPS-199 {baseline} Baseline  ·  System Security Plan",
        11,
        False,
        RGBColor(0x6B, 0x72, 0x80),
    )
    doc.add_paragraph()

    meta = [
        ("Organization", project.get("customer_name") or ""),
        ("Information System", project.get("system_name") or ""),
        ("Security Categorization / Baseline", f"FIPS-199 {baseline}"),
        ("Cloud Platform", project.get("platform") or ""),
        ("Document Version", project.get("version") or "0.1"),
        ("Date", project.get("document_date") or ""),
        ("Prepared By", project.get("prepared_by") or ""),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = 1
    for i, (label, value) in enumerate(meta):
        lc = table.cell(i, 0)
        _shade(lc, LABEL_FILL)
        _runs(lc, label, bold=True)
        _runs(table.cell(i, 1), str(value))
    doc.add_page_break()


def _front_matter(doc: Any, project: Mapping[str, Any]) -> None:
    doc.add_heading("How to Use This Document", level=1)
    baseline = _baseline_label(project)
    doc.add_paragraph(
        f"This System Security Plan documents implementation of the NIST SP 800-53 Rev 5 "
        f"{baseline} baseline control set (NIST SP 800-53B). Narratives are draft "
        "assessor-facing language scaffolded from the catalog; tailor each entry to the "
        "actual environment, and fill every organization-defined parameter (ODP) before "
        "authorization."
    )
    doc.add_heading("Assumptions and Tailoring Notes", level=2)
    for note in (
        "Control statements are drafts and must be reviewed and confirmed against the "
        "actual system before assessment.",
        "Organization-Defined Parameters (ODPs) shown as [unset] must be filled with the "
        "organization's actual assigned value.",
        "Implementation Status must be confirmed (Implemented, Partially Implemented, "
        "Planned, Alternative Implementation, or Not Applicable).",
        "Control Origination follows FedRAMP-style concepts (service provider corporate, "
        "system-specific, hybrid, inherited, customer-configured).",
        "Where a cloud or external service provider is used, inheritance must be confirmed "
        "against that provider's authorization and customer responsibility matrix.",
    ):
        doc.add_paragraph(note, style="List Bullet")
    doc.add_page_break()


def _odp_lines(entry: Mapping[str, Any]) -> list[str]:
    odps: Mapping[str, Any] = entry.get("odp_values") or {}
    lines = []
    for param_id, value in odps.items():
        shown = str(value).strip() if value not in (None, "") else "[unset]"
        lines.append(f"{param_id}: {shown}")
    return lines


def _add_control(doc: Any, entry: Mapping[str, Any]) -> None:
    cid = entry.get("control_id") or ""
    title = entry.get("title") or ""
    h = doc.add_heading(level=2)
    run = h.add_run(f"{cid} — {title}")
    run.font.color.rgb = ACCENT

    rows = [
        f"{cid} — {title}",
        "Statement:  "
        + " ".join(str(p.get("text") or "") for p in (entry.get("part_narratives") or [])),
        "Organization-Defined Parameters:  " + ("; ".join(_odp_lines(entry)) or "—"),
        f"Responsible Role:  {entry.get('responsible_role') or '—'}",
        "Implementation Status:  " + (", ".join(entry.get("implementation_status") or []) or "—"),
        "Control Origination:  " + (", ".join(entry.get("control_origination") or []) or "—"),
    ]
    table = doc.add_table(rows=len(rows), cols=1)
    table.style = "Table Grid"
    for i, text in enumerate(rows):
        cell = table.cell(i, 0)
        if i == 0:
            _shade(cell, HEADER_FILL)
            _runs(cell, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        else:
            _shade(cell, LABEL_FILL if i in (2, 3) else "FFFFFF")
            _runs(cell, text)
    doc.add_paragraph()


def render_80053_docx(project: Mapping[str, Any], entries: Sequence[Mapping[str, Any]]) -> bytes:
    """Build the NIST SP 800-53 Rev 5 SSP and return the ``.docx`` file as bytes."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _cover(doc, project)
    _front_matter(doc, project)

    by_family: dict[str, list[Mapping[str, Any]]] = {}
    for e in entries:
        by_family.setdefault(e.get("domain") or "?", []).append(e)

    ordered = [f for f in FAMILY_ORDER if f in by_family]
    ordered += [f for f in by_family if f not in FAMILY_ORDER]

    for family in ordered:
        doc.add_heading(f"{family} — {_family_name(family)}", level=1)
        controls = sorted(by_family[family], key=_sort_key)
        for entry in controls:
            _add_control(doc, entry)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
