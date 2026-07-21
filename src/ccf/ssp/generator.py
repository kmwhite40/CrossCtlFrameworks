"""Render a CMMC Level 2 / NIST SP 800-171 Rev.2 SSP project to ``.docx``.

The document covers the 110 CMMC Level 2 practices (NIST SP 800-171 Rev.2) —
it is not a FedRAMP SSP and does not contain NIST SP 800-53 content. The
layout format resembles a FedRAMP-style Appendix A (the source template was
``CMMC_L2_SSP_Appendix_A_NIST_800-171_FedRAMP_Style.docx``): a cover/metadata
block, usage notes, control-origination definitions, then one section per
CMMC domain containing, for each control, a *Control Summary Information*
table and a *What is the solution and how is it implemented?* table with one
row per NIST 800-171A determination part.
"""

from __future__ import annotations

import io
from collections.abc import Mapping, Sequence
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from . import constants

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)  # deep navy
HEADER_FILL = "1F3A5F"
LABEL_FILL = "E8EDF3"
CHECKED, UNCHECKED = "☒", "☐"


def _shade(cell: Any, fill: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _runs(cell: Any, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color is not None:
        run.font.color.rgb = color


def _checkbox_line(options: Sequence[str], selected: Sequence[str]) -> str:
    chosen = {s.strip().lower() for s in selected}
    return "    ".join(
        f"{CHECKED if opt.lower() in chosen else UNCHECKED} {opt}" for opt in options
    )


def _add_summary_table(doc: Any, entry: Mapping[str, Any]) -> None:
    cid, nist = entry["control_id"], entry.get("nist_id") or ""
    rows = [
        f"{cid} — {nist} Control Summary Information",
        f"CMMC L2 Practice / NIST 800-171 Requirement:  {cid}  ({nist})",
        f"Responsible Role:  {entry.get('responsible_role') or '—'}",
        "Implementation Status (check all that apply):  "
        + _checkbox_line(
            constants.IMPLEMENTATION_STATUS_OPTIONS, entry.get("implementation_status") or []
        ),
        "Control Origination (check all that apply):  "
        + _checkbox_line(
            constants.CONTROL_ORIGINATION_OPTIONS, entry.get("control_origination") or []
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=1)
    table.style = "Table Grid"
    for i, text in enumerate(rows):
        cell = table.cell(i, 0)
        if i == 0:
            _shade(cell, HEADER_FILL)
            _runs(cell, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        else:
            _shade(cell, LABEL_FILL if i in (1, 2) else "FFFFFF")
            _runs(cell, text)


def _add_solution_table(doc: Any, entry: Mapping[str, Any]) -> None:
    cid, nist = entry["control_id"], entry.get("nist_id") or ""
    parts: list[Mapping[str, str]] = list(entry.get("part_narratives") or [])
    table = doc.add_table(rows=len(parts) + 1, cols=2)
    table.style = "Table Grid"
    header = table.cell(0, 0).merge(table.cell(0, 1))
    _shade(header, HEADER_FILL)
    _runs(
        header,
        f"{cid} — {nist} What is the solution and how is it implemented?",
        bold=True,
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )
    for i, part in enumerate(parts, start=1):
        label = part.get("label", "")
        left = table.cell(i, 0)
        _shade(left, LABEL_FILL)
        _runs(left, f"Part [{label}]" if label else "Statement", bold=True)
        _runs(table.cell(i, 1), part.get("text", ""))
    # Narrow the label column.
    for row in table.rows:
        row.cells[0].width = Pt(70)


def _add_control(doc: Any, entry: Mapping[str, Any]) -> None:
    cid = entry["control_id"]
    nist = entry.get("nist_id") or ""
    title = entry.get("title") or ""
    h = doc.add_heading(level=2)
    run = h.add_run(f"{cid}  (NIST SP 800-171: {nist}) — {title}")
    run.font.color.rgb = ACCENT
    meta = doc.add_paragraph()
    mr = meta.add_run(
        f"CMMC 2.0 Level 2 practice {cid}  |  NIST SP 800-171 Rev. 2 requirement {nist}"
    )
    mr.italic = True
    mr.font.size = Pt(8.5)
    mr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    _add_summary_table(doc, entry)
    doc.add_paragraph()
    _add_solution_table(doc, entry)
    doc.add_paragraph()


def _cover(doc: Any, project: Mapping[str, Any]) -> None:
    def centered(text: str, size: int, bold: bool, color: RGBColor | None = None) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color

    for _ in range(3):
        doc.add_paragraph()
    centered(project.get("title") or "System Security Plan (SSP)", 26, True, ACCENT)
    centered("Appendix A — CMMC Level 2 Security Requirements", 15, True)
    centered(
        "NIST SP 800-171 Rev. 2  ·  System Security Plan, documented in an Appendix-A layout",
        11,
        False,
        RGBColor(0x6B, 0x72, 0x80),
    )
    doc.add_paragraph()

    meta = [
        ("Organization", project.get("customer_name") or ""),
        ("Information System", project.get("system_name") or ""),
        ("CMMC Model / Level", "CMMC 2.0 — Level 2 (NIST SP 800-171 Rev. 2, 110 requirements)"),
        ("Cloud Platform", project.get("platform") or ""),
        ("Document Version", project.get("version") or "0.1"),
        ("Date", project.get("document_date") or ""),
        ("Prepared By", project.get("prepared_by") or ""),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = 1
    for i, (label, value) in enumerate(meta):
        lc = table.cell(i, 0)
        _shade(lc, LABEL_FILL)
        _runs(lc, label, bold=True)
        _runs(table.cell(i, 1), str(value))
    doc.add_page_break()


def _front_matter(doc: Any) -> None:
    doc.add_heading("How to Use This Document", level=1)
    doc.add_paragraph(
        "This appendix documents implementation of the 110 CMMC Level 2 security requirements "
        "(NIST SP 800-171 Rev. 2). All narratives are drafts written in assessor-facing "
        "language; tailor each entry to the actual environment before assessment."
    )
    doc.add_heading("Assumptions and Tailoring Notes", level=2)
    for note in (
        "The assessment boundary includes the systems, users, devices, cloud services, "
        "security tooling, and facilities that store, process, or transmit CUI.",
        "Implementation Status is shown as a draft target and must be confirmed "
        "(Implemented, Partially Implemented, Planned, Alternative Implementation, "
        "or Not Applicable).",
        "Control Origination uses FedRAMP-style concepts adapted for CMMC.",
        "Where a cloud or external service provider is used, inheritance must be confirmed "
        "against that provider's authorization and customer responsibility matrix.",
        "Part labels [a], [b], [c]… correspond to the determination statements in "
        "NIST SP 800-171A.",
    ):
        doc.add_paragraph(note, style="List Bullet")

    doc.add_heading("Control Origination Definitions", level=1)
    table = doc.add_table(rows=len(constants.ORIGINATION_DEFINITIONS) + 1, cols=2)
    table.style = "Table Grid"
    _shade(table.cell(0, 0), HEADER_FILL)
    _shade(table.cell(0, 1), HEADER_FILL)
    _runs(table.cell(0, 0), "Control Origination", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _runs(
        table.cell(0, 1),
        "Definition for this CMMC SSP",
        bold=True,
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )
    for i, (name, definition) in enumerate(constants.ORIGINATION_DEFINITIONS, start=1):
        _shade(table.cell(i, 0), LABEL_FILL)
        _runs(table.cell(i, 0), name, bold=True)
        _runs(table.cell(i, 1), definition)
    doc.add_page_break()


def _kv_table(doc: Any, rows: Sequence[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lc = table.cell(i, 0)
        _shade(lc, LABEL_FILL)
        _runs(lc, label, bold=True)
        _runs(table.cell(i, 1), str(value) if value else "—")


def _system_characterization(doc: Any, project: Mapping[str, Any]) -> None:
    """Front matter (FedRAMP-Appendix-A-style layout): categorization, roles,
    boundary, interconnections — content is CMMC L2 / NIST 800-171, not 800-53.
    """
    meta: Mapping[str, Any] = project.get("metadata_json") or {}
    fips = meta.get("fips199") or {}
    roles = meta.get("roles") or {}

    doc.add_heading("1. System Characterization", level=1)
    _kv_table(
        doc,
        [
            ("Information System Name", project.get("system_name") or ""),
            ("System Type", meta.get("system_type") or ""),
            ("Cloud Platform / Environment", project.get("platform") or ""),
            ("Operational Status", meta.get("operational_status") or "Operational"),
        ],
    )

    doc.add_heading("1.1 Security Categorization (FIPS-199)", level=2)
    _kv_table(
        doc,
        [
            ("Confidentiality", fips.get("confidentiality") or ""),
            ("Integrity", fips.get("integrity") or ""),
            ("Availability", fips.get("availability") or ""),
            ("Overall Categorization", fips.get("overall") or ""),
        ],
    )

    doc.add_heading("1.2 Roles and Responsibilities", level=2)

    def _role(key: str) -> str:
        r = roles.get(key) or {}
        name, email = r.get("name") or "", r.get("email") or ""
        return f"{name}{'  ·  ' + email if email else ''}" if name else ""

    _kv_table(
        doc,
        [
            ("System Owner", _role("system_owner")),
            ("Information System Security Officer (ISSO)", _role("isso")),
            ("Information System Security Manager (ISSM)", _role("issm")),
            ("Authorizing Official (AO)", _role("authorizing_official")),
        ],
    )

    doc.add_heading("1.3 Authorization Boundary", level=2)
    doc.add_paragraph(
        meta.get("authorization_boundary")
        or "The authorization boundary includes the systems, users, devices, cloud services, "
        "security tooling, and facilities that store, process, or transmit CUI."
    )

    leveraged = meta.get("leveraged_authorizations") or []
    if leveraged:
        doc.add_heading("1.4 Leveraged Authorizations (Inheritance)", level=2)
        table = doc.add_table(rows=len(leveraged) + 1, cols=2)
        table.style = "Table Grid"
        _shade(table.cell(0, 0), HEADER_FILL)
        _shade(table.cell(0, 1), HEADER_FILL)
        _runs(table.cell(0, 0), "Provider / Service", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _runs(table.cell(0, 1), "Authorization", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        for i, lv in enumerate(leveraged, start=1):
            _runs(table.cell(i, 0), str(lv.get("provider") or ""))
            _runs(table.cell(i, 1), str(lv.get("authorization") or ""))

    laws = meta.get("laws_regulations") or []
    if laws:
        doc.add_heading("1.5 Applicable Laws and Regulations", level=2)
        for law in laws:
            doc.add_paragraph(str(law), style="List Bullet")

    doc.add_page_break()


def _revision_history(doc: Any, project: Mapping[str, Any]) -> None:
    history = project.get("revision_history") or []
    if not history:
        return
    doc.add_heading("Document Revision History", level=1)
    table = doc.add_table(rows=len(history) + 1, cols=4)
    table.style = "Table Grid"
    for j, head in enumerate(("Version", "Date", "Author", "Description")):
        _shade(table.cell(0, j), HEADER_FILL)
        _runs(table.cell(0, j), head, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, rev in enumerate(history, start=1):
        _runs(table.cell(i, 0), str(rev.get("version") or ""))
        _runs(table.cell(i, 1), str(rev.get("date") or ""))
        _runs(table.cell(i, 2), str(rev.get("author") or ""))
        _runs(table.cell(i, 3), str(rev.get("notes") or ""))
    doc.add_page_break()


def generate_ssp_docx(project: Mapping[str, Any], entries: Sequence[Mapping[str, Any]]) -> bytes:
    """Build the SSP and return the ``.docx`` file as bytes."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _cover(doc, project)
    _revision_history(doc, project)
    _system_characterization(doc, project)
    _front_matter(doc)

    by_domain: dict[str, list[Mapping[str, Any]]] = {}
    for e in entries:
        by_domain.setdefault(e.get("domain") or "?", []).append(e)

    ordered = [d for d in constants.DOMAIN_ORDER if d in by_domain]
    ordered += [d for d in by_domain if d not in constants.DOMAIN_ORDER]

    for domain in ordered:
        section = constants.section_number(domain)
        name = constants.domain_name(domain)
        doc.add_heading(f"{section} {name} ({domain})", level=1)
        controls = sorted(by_domain[domain], key=lambda x: x.get("control_id") or "")
        for entry in controls:
            _add_control(doc, entry)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
