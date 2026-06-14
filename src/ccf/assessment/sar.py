"""Render a CMMC L2 Security Assessment Report (SAR) to a .docx."""

from __future__ import annotations

import io
from collections.abc import Mapping, Sequence
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..ssp import constants

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
HEADER_FILL = "1F3A5F"
LABEL_FILL = "E8EDF3"

_FINDING_LABEL = {
    "satisfied": "Satisfied",
    "other_than_satisfied": "Other Than Satisfied",
    "not_applicable": "Not Applicable",
    "not_assessed": "Not Assessed",
}
_FINDING_FILL = {
    "satisfied": "E3F4E9",
    "other_than_satisfied": "FBE3E4",
    "not_applicable": "ECEFF3",
    "not_assessed": "FFF6E3",
}


def _shade(cell: Any, fill: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _runs(cell: Any, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color is not None:
        run.font.color.rgb = color


def _centered(doc: Any, text: str, size: int, bold: bool, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def _cover(doc: Any, meta: Mapping[str, Any]) -> None:
    for _ in range(3):
        doc.add_paragraph()
    _centered(doc, "SECURITY ASSESSMENT REPORT (SAR)", 26, True, ACCENT)
    _centered(doc, "CMMC Level 2 — NIST SP 800-171 Rev. 2 (110 practices)", 14, True)
    _centered(
        doc,
        "Examine · Interview · Test — determination of satisfied / other than satisfied",
        11,
        False,
        RGBColor(0x6B, 0x72, 0x80),
    )
    doc.add_paragraph()
    rows = [
        ("Organization", meta.get("customer_name") or ""),
        ("Information System", meta.get("system_name") or ""),
        ("Assessment", meta.get("assessment_name") or ""),
        ("Assessment Type", meta.get("kind") or ""),
        ("Lead Assessor", meta.get("assessor") or ""),
        ("Assessment Period", meta.get("period") or ""),
        ("Report Date", meta.get("date") or ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = 1
    for i, (label, value) in enumerate(rows):
        lc = table.cell(i, 0)
        _shade(lc, LABEL_FILL)
        _runs(lc, label, bold=True)
        _runs(table.cell(i, 1), str(value))
    doc.add_page_break()


def _summary(doc: Any, summary: Mapping[str, Any]) -> None:
    doc.add_heading("1. Assessment Methodology", level=1)
    doc.add_paragraph(
        "Each CMMC Level 2 practice was assessed against its NIST SP 800-171A determination "
        "statements using the examine, interview, and test methods. A practice is Satisfied "
        "only when every determination statement is met; any unmet statement results in an "
        "Other Than Satisfied finding and a corresponding POA&M item."
    )

    doc.add_heading("2. Summary of Findings", level=1)
    counts = summary.get("by_finding", {})
    t = doc.add_table(rows=2, cols=4)
    t.style = "Table Grid"
    heads = ["Satisfied", "Other Than Satisfied", "Not Applicable", "Not Assessed"]
    keys = ["satisfied", "other_than_satisfied", "not_applicable", "not_assessed"]
    for j, h in enumerate(heads):
        _shade(t.cell(0, j), HEADER_FILL)
        _runs(t.cell(0, j), h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _runs(t.cell(1, j), str(counts.get(keys[j], 0)))
    doc.add_paragraph(
        f"Practices assessed: {summary.get('assessed', 0)} of {summary.get('total', 0)} "
        f"({summary.get('complete_pct', 0)}%). Reviewer-signed: {summary.get('reviewed', 0)}."
    )

    by_domain = summary.get("by_domain", {})
    if by_domain:
        doc.add_heading("2.1 Findings by Domain", level=2)
        dt = doc.add_table(rows=len(by_domain) + 1, cols=4)
        dt.style = "Table Grid"
        for j, h in enumerate(["Domain", "Practices", "Satisfied", "Other Than Satisfied"]):
            _shade(dt.cell(0, j), HEADER_FILL)
            _runs(dt.cell(0, j), h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        ordered = [d for d in constants.DOMAIN_ORDER if d in by_domain]
        ordered += [d for d in by_domain if d not in constants.DOMAIN_ORDER]
        for i, d in enumerate(ordered, start=1):
            b = by_domain[d]
            _shade(dt.cell(i, 0), LABEL_FILL)
            _runs(dt.cell(i, 0), f"{d} — {constants.domain_name(d)}", bold=True)
            _runs(dt.cell(i, 1), str(b.get("total", 0)))
            _runs(dt.cell(i, 2), str(b.get("satisfied", 0)))
            _runs(dt.cell(i, 3), str(b.get("other_than_satisfied", 0)))
    doc.add_page_break()


def _detail(doc: Any, results: Sequence[Mapping[str, Any]]) -> None:
    doc.add_heading("3. Detailed Results", level=1)
    by_domain: dict[str, list[Mapping[str, Any]]] = {}
    for r in results:
        by_domain.setdefault(r.get("domain") or "?", []).append(r)
    ordered = [d for d in constants.DOMAIN_ORDER if d in by_domain]
    ordered += [d for d in by_domain if d not in constants.DOMAIN_ORDER]

    for d in ordered:
        doc.add_heading(f"{constants.section_number(d)} {constants.domain_name(d)} ({d})", level=2)
        for r in sorted(by_domain[d], key=lambda x: x.get("control_id") or ""):
            finding = r.get("finding", "not_assessed")
            h = doc.add_heading(level=3)
            run = h.add_run(
                f"{r.get('control_id')}  ({r.get('nist_id')}) — {r.get('title') or ''}"
            )
            run.font.color.rgb = ACCENT

            obj = r.get("objective_findings") or []
            table = doc.add_table(rows=4 + len(obj), cols=2)
            table.style = "Table Grid"
            cells = [
                ("Finding", _FINDING_LABEL.get(finding, finding)),
                ("Examine", r.get("examine_note") or "—"),
                ("Interview", r.get("interview_note") or "—"),
                ("Test", r.get("test_note") or "—"),
            ]
            for i, (label, value) in enumerate(cells):
                lc = table.cell(i, 0)
                _shade(lc, _FINDING_FILL.get(finding, LABEL_FILL) if i == 0 else LABEL_FILL)
                _runs(lc, label, bold=True)
                _runs(table.cell(i, 1), str(value))
            for k, part in enumerate(obj, start=4):
                lc = table.cell(k, 0)
                _shade(lc, LABEL_FILL)
                lbl = part.get("label", "")
                _runs(lc, f"Part [{lbl}]" if lbl else "Statement", bold=True)
                pf = _FINDING_LABEL.get(part.get("finding", "not_assessed"), "")
                _runs(table.cell(k, 1), f"[{pf}] {part.get('text', '')}".strip())

            extra: list[str] = []
            if r.get("assessor_note"):
                extra.append(f"Assessor note: {r['assessor_note']}")
            if r.get("evidence_ref"):
                extra.append(f"Evidence: {r['evidence_ref']}")
            if extra:
                note = doc.add_paragraph()
                nr = note.add_run("  ".join(extra))
                nr.italic = True
                nr.font.size = Pt(8.5)
                nr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            doc.add_paragraph()


def generate_sar_docx(
    meta: Mapping[str, Any],
    summary: Mapping[str, Any],
    results: Sequence[Mapping[str, Any]],
) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    _cover(doc, meta)
    _summary(doc, summary)
    _detail(doc, results)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
