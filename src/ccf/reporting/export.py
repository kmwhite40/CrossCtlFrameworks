"""Render a custom compliance report to .xlsx or .docx bytes.

Both renderers consume the same ``(summary, rows)`` shape produced by the
report builder in :mod:`ccf.api.routes.reports`, so the API can offer CSV,
JSON, XLSX, and DOCX from one query.
"""

from __future__ import annotations

import io
from collections.abc import Mapping, Sequence
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

# Column order / human labels for the tabular report rows.
_COLUMNS: tuple[tuple[str, str], ...] = (
    ("identifier", "Control"),
    ("family", "Family"),
    ("control_name", "Control Name"),
    ("baseline_low", "Low"),
    ("baseline_mod", "Moderate"),
    ("baseline_high", "High"),
    ("implementation_status", "Implementation"),
    ("responsibility", "Responsibility"),
    ("owner", "Owner"),
    ("last_assessed_on", "Last Assessed"),
    ("crosswalk_framework", "Crosswalk"),
    ("crosswalk_values", "Crosswalk Values"),
    ("ai_sourced", "AI-Sourced"),
)

_BRAND = "4F46E5"  # indigo, matches the app brand
_HEADER_FILL = PatternFill("solid", fgColor=_BRAND)


def _cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else ""
    return str(value)


def _summary_pairs(summary: Mapping[str, Any]) -> list[tuple[str, str]]:
    return [
        ("Generated", _cell(summary.get("generated_at"))),
        ("Organization", _cell(summary.get("organization")) or "(catalog)"),
        ("System", _cell(summary.get("system")) or "(any)"),
        ("Baseline", _cell(summary.get("baseline")) or "(all)"),
        ("Crosswalk framework", _cell(summary.get("framework")) or "(none)"),
        ("Family filter", _cell(summary.get("family_filter")) or "(all)"),
        ("Total rows", _cell(summary.get("total_rows"))),
    ]


def _risk_summary_pairs(risk_summary: Mapping[str, Any] | None) -> list[tuple[str, str]]:
    """POA&M/risk posture rows for the report, reconciled 1:1 to the dashboard's
    ``ccf.analytics.posture.org_summary`` for the same scope (CISO-10) — same
    numbers, not a re-derivation, so the export can never silently drift from
    what leadership sees on screen."""
    if not risk_summary:
        return []
    worst = risk_summary.get("worst_system") or {}
    risks_by_status = risk_summary.get("risks_by_status") or {}
    open_risks = sum(n for status, n in risks_by_status.items() if status != "closed")
    return [
        ("Systems scored", _cell(risk_summary.get("systems_scored"))),
        ("Avg SPRS score", _cell(risk_summary.get("avg_sprs_score"))),
        ("Min (worst) SPRS score", _cell(risk_summary.get("min_sprs_score"))),
        ("Worst-scoring system", _cell(worst.get("name")) or "(none)"),
        ("Open POA&Ms", _cell(risk_summary.get("open_poams"))),
        ("Overdue POA&Ms", _cell(risk_summary.get("overdue_poams"))),
        ("Risk-accepted POA&Ms", _cell(risk_summary.get("accepted_poams"))),
        ("Open risks", _cell(open_risks)),
    ]


def report_to_xlsx(
    summary: Mapping[str, Any],
    rows: Sequence[Mapping[str, Any]],
    risk_summary: Mapping[str, Any] | None = None,
) -> bytes:
    """Render the report as a two-sheet workbook (Summary + Controls).

    ``risk_summary`` (CISO-10) adds a POA&M/risk posture block beneath the
    report metadata on the Summary sheet, reconciled to the dashboard.
    """
    wb = Workbook()

    meta = wb.active
    meta.title = "Summary"
    meta["A1"] = "Concord compliance report"
    meta["A1"].font = Font(size=14, bold=True, color=_BRAND)
    row_i = 3
    for label, value in _summary_pairs(summary):
        meta.cell(row=row_i, column=1, value=label).font = Font(bold=True)
        meta.cell(row=row_i, column=2, value=value)
        row_i += 1
    risk_pairs = _risk_summary_pairs(risk_summary)
    if risk_pairs:
        row_i += 1
        meta.cell(row=row_i, column=1, value="POA&M / Risk Posture").font = Font(
            bold=True, color=_BRAND
        )
        row_i += 1
        for label, value in risk_pairs:
            meta.cell(row=row_i, column=1, value=label).font = Font(bold=True)
            meta.cell(row=row_i, column=2, value=value)
            row_i += 1
    meta.column_dimensions["A"].width = 22
    meta.column_dimensions["B"].width = 48

    ws = wb.create_sheet("Controls")
    headers = [label for _, label in _COLUMNS]
    ws.append(headers)
    for col, _ in enumerate(headers, start=1):
        c = ws.cell(row=1, column=col)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = _HEADER_FILL
        c.alignment = Alignment(vertical="center", wrap_text=True)
    for row in rows:
        ws.append([_cell(row.get(key)) for key, _ in _COLUMNS])

    # Reasonable widths + a frozen, filterable header.
    widths = (16, 9, 40, 7, 10, 7, 16, 16, 12, 14, 12, 32, 12)
    for col, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws.max_row}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _shade(cell: Any, fill: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {qn("w:fill"): fill})
    cell._tc.get_or_add_tcPr().append(shd)


def report_to_docx(
    summary: Mapping[str, Any],
    rows: Sequence[Mapping[str, Any]],
    risk_summary: Mapping[str, Any] | None = None,
) -> bytes:
    """Render the report as a titled Word document with a control table.

    ``risk_summary`` (CISO-10) adds a POA&M/risk posture section, reconciled
    to the dashboard, before the control table.
    """
    doc = Document()

    title = doc.add_heading(level=0)
    run = title.add_run("Concord Compliance Report")
    run.font.color.rgb = RGBColor.from_string(_BRAND)

    for label, value in _summary_pairs(summary):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    risk_pairs = _risk_summary_pairs(risk_summary)
    if risk_pairs:
        doc.add_heading("POA&M / Risk Posture", level=1)
        for label, value in risk_pairs:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)

    doc.add_heading("Controls", level=1)
    headers = [label for _, label in _COLUMNS]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        _shade(cell, _BRAND)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, (key, _) in enumerate(_COLUMNS):
            cells[i].text = _cell(row.get(key))
            for para in cells[i].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
